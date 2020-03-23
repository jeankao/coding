# -*- coding: UTF-8 -*-
from django.shortcuts import render, HttpResponseRedirect
from django.shortcuts import render_to_response, redirect
from django.template import RequestContext
from django.views.generic import ListView, DetailView, CreateView, UpdateView, RedirectView
from teacher.models import Classroom, ImportUser, TWork, Assistant, CWork, FWork, FClass, FContent
from student.models import *
from account.models import Message, MessagePoll, MessageContent, PointHistory
from account.avatar import *
from teacher.forms import *
from django.contrib.auth.decorators import login_required
from student.lesson import *
from django.core.exceptions import ObjectDoesNotExist, MultipleObjectsReturned
from django.contrib.auth.models import User
from django.utils import timezone
from django.contrib.auth.decorators import login_required, user_passes_test
from django.http import JsonResponse
from django.core.exceptions import PermissionDenied
from django.utils.decorators import method_decorator
from django.db.models import Q
from django.forms.models import model_to_dict
from django.core.files.storage import FileSystemStorage
from wsgiref.util import FileWrapper
from collections import OrderedDict
import django_excel as excel
from account.forms import PasswordForm, RealnameForm
import sys
from uuid import uuid4
import os
import StringIO
from shutil import copyfile
import xlsxwriter
from datetime import datetime
from django.http import HttpResponse
from django.utils.timezone import localtime
import json
from docx import *
from docx.shared import Inches
from docx.shared import RGBColor
from docx.oxml.shared import OxmlElement, qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from datetime import datetime
from helper import VideoLogHelper
import re
import io
reload(sys)
from django.db.models.functions import Length
from django.forms import modelformset_factory
from PIL import Image

sys.setdefaultencoding('utf-8')

def filename_browser(request, filename):
	browser = request.META['HTTP_USER_AGENT'].lower()
	if 'edge' in browser:
		response['Content-Disposition'] = 'attachment; filename='+urlquote(filename)+'; filename*=UTF-8\'\'' + urlquote(filename)
		return response			
	elif 'webkit' in browser:
		# Safari 3.0 and Chrome 2.0 accepts UTF-8 encoded string directly.
		filename_header = 'filename=%s' % filename.encode('utf-8').decode('ISO-8859-1')
	elif 'trident' in browser or 'msie' in browser:
		# IE does not support internationalized filename at all.
		# It can only recognize internationalized URL, so we do the trick via routing rules.
		filename_header = 'filename='+filename.encode("BIG5").decode("ISO-8859-1")					
	else:
		# For others like Firefox, we follow RFC2231 (encoding extension in HTTP headers).
		filename_header = 'filename*="utf8\'\'' + str(filename.encode('utf-8').decode('ISO-8859-1')) + '"'
	return filename_header		


def is_assistant(user, classroom_id):
    assistants = Assistant.objects.filter(classroom_id=classroom_id, user_id=user.id)
    if len(assistants)>0 :
        return True
    return False

# 判斷是否為授課教師
def is_teacher(user, classroom_id):
    return user.groups.filter(name='teacher').exists() and Classroom.objects.filter(teacher_id=user.id, id=classroom_id).exists()

def not_in_teacher_group(user):
    if not user.groups.filter(name='teacher').exists():
        if not Assistant.objects.filter(user_id=user.id).exists():
            return False
    return True

# 列出所有課程
class ClassroomListView(ListView):
    model = Classroom
    context_object_name = 'classrooms'
    paginate_by = 30
    def dispatch(self, *args, **kwargs):
        if not not_in_teacher_group(self.request.user):
            raise PermissionDenied
        else :
            return super(ClassroomListView, self).dispatch(*args, **kwargs)
    def get_queryset(self):
        queryset = Classroom.objects.filter(teacher_id=self.request.user.id).order_by("-id")
        return queryset

#新增一個課程
class ClassroomCreateView(CreateView):
    model = Classroom
    form_class = ClassroomForm
    template_name = 'form.html'
    def dispatch(self, *args, **kwargs):
        if not not_in_teacher_group(self.request.user):
            raise PermissionDenied
        else :
            return super(ClassroomCreateView, self).dispatch(*args, **kwargs)
    def form_valid(self, form):
        self.object = form.save(commit=False)
        self.object.teacher_id = self.request.user.id
        self.object.save()
        # 將教師設為0號學生
        enroll = Enroll(classroom_id=self.object.id, student_id=self.request.user.id, seat=0)
        enroll.save()

        return redirect("/teacher/classroom")

# 修改選課密碼
@login_required
@user_passes_test(not_in_teacher_group, login_url='/')
def classroom_edit(request, classroom_id):
    classroom = Classroom.objects.get(id=classroom_id)
    if request.method == 'POST':
        form = ClassroomForm(request.POST)
        if form.is_valid():
            classroom.name =form.cleaned_data['name']
            classroom.password = form.cleaned_data['password']
            classroom.save()
            return redirect('/teacher/classroom')
    else:
        form = ClassroomForm(instance=classroom)

    return render(request, 'teacher/classroom_form.html',{'form': form})

# 退選
@login_required
@user_passes_test(not_in_teacher_group, login_url='/')
def unenroll(request, enroll_id, classroom_id):
    enroll = Enroll.objects.get(id=enroll_id)
    enroll.delete()
    classroom_name = Classroom.objects.get(id=classroom_id).name

    return redirect('/student/classmate/'+classroom_id)

# 設定班級助教
@login_required
@user_passes_test(not_in_teacher_group, login_url='/')
def classroom_assistant(request, classroom_id):
    assistants = Assistant.objects.filter(classroom_id=classroom_id).order_by("-id")
    classroom = Classroom.objects.get(id=classroom_id)

    return render(request, 'teacher/assistant.html',{'assistants': assistants, 'classroom':classroom})

# 教師可以查看所有帳號
class AssistantListView(ListView):
    context_object_name = 'users'
    paginate_by = 20
    template_name = 'teacher/assistant_user.html'
    def dispatch(self, *args, **kwargs):
        if not not_in_teacher_group(self.request.user):
            raise PermissionDenied
        else :
            return super(AssistantListView, self).dispatch(*args, **kwargs)
    def get_queryset(self):
        if self.request.GET.get('account') != None:
            keyword = self.request.GET.get('account')
            queryset = User.objects.filter(Q(username__icontains=keyword) | Q(first_name__icontains=keyword)).order_by('-id')
        else :
            queryset = User.objects.all().order_by('-id')
        return queryset

    def get_context_data(self, **kwargs):
        context = super(AssistantListView, self).get_context_data(**kwargs)
        context['classroom'] = Classroom.objects.get(id=self.kwargs['classroom_id'])
        assistant_list = []
        assistants = Assistant.objects.filter(classroom_id=self.kwargs['classroom_id'])
        for assistant in assistants:
            assistant_list.append(assistant.user_id)
        context['assistants'] = assistant_list
        return context

# 列出所有助教課程
class AssistantClassroomListView(ListView):
    model = Classroom
    context_object_name = 'classrooms'
    template_name = 'teacher/assistant_list.html'
    paginate_by = 20
    def dispatch(self, *args, **kwargs):
        if not not_in_teacher_group(self.request.user):
            raise PermissionDenied
        else :
            return super(AssistantClassroomListView, self).dispatch(*args, **kwargs)
    def get_queryset(self):
        assistants = Assistant.objects.filter(user_id=self.request.user.id)
        classroom_list = []
        for assistant in assistants:
            classroom_list.append(assistant.classroom_id)
        queryset = Classroom.objects.filter(id__in=classroom_list).order_by("-id")
        return queryset

# Ajax 設為助教、取消助教
def assistant_make(request):
    classroom_id = request.POST.get('classroomid')
    user_id = request.POST.get('userid')
    action = request.POST.get('action')
    if user_id and action :
        if action == 'set':
            try :
                assistant = Assistant.objects.get(classroom_id=classroom_id, user_id=user_id)
            except ObjectDoesNotExist :
                assistant = Assistant(classroom_id=classroom_id, user_id=user_id)
                assistant.save()
            # 將教師設為0號學生
            enroll = Enroll(classroom_id=classroom_id, student_id=user_id, seat=0)
            enroll.save()
        else :
            try :
                assistant = Assistant.objects.get(classroom_id=classroom_id, user_id=user_id)
                assistant.delete()
                enroll = Enroll.objects.filter(classroom_id=classroom_id, student_id=user_id)
                enroll.delete()
            except ObjectDoesNotExist :
                pass
        return JsonResponse({'status':'ok'}, safe=False)
    else:
        return JsonResponse({'status':'fail'}, safe=False)

# 退選
@login_required
@user_passes_test(not_in_teacher_group, login_url='/')
def unenroll(request, enroll_id, classroom_id):
    enroll = Enroll.objects.get(id=enroll_id)
    enroll.delete()
    classroom_name = Classroom.objects.get(id=classroom_id).name
    return redirect('/student/classmate/'+classroom_id)

# 列出所有公告
class AnnounceListView(ListView):
    model = Message
    context_object_name = 'messages'
    template_name = 'teacher/announce_list.html'
    paginate_by = 20
    def dispatch(self, *args, **kwargs):
        if not not_in_teacher_group(self.request.user):
            raise PermissionDenied
        else :
            return super(AnnounceListView, self).dispatch(*args, **kwargs)
    def get_queryset(self):
        queryset = Message.objects.filter(classroom_id=self.kwargs['classroom_id'], author_id=self.request.user.id).order_by("-id")
        return queryset

    def get_context_data(self, **kwargs):
        context = super(AnnounceListView, self).get_context_data(**kwargs)
        context['classroom'] = Classroom.objects.get(id=self.kwargs['classroom_id'])
        return context

    # 限本班任課教師
    def render(request, self, context):
        if not is_teacher(self.request.user, self.kwargs['classroom_id']) and not is_assistant(self.request.user, self.kwargs['classroom_id']):
            return redirect('/')
        return super(AnnounceListView, self).render(request, context)

#新增一個公告
class AnnounceCreateView(CreateView):
    model = Message
    form_class = AnnounceForm
    template_name = 'teacher/announce_form.html'
    def dispatch(self, *args, **kwargs):
        if not not_in_teacher_group(self.request.user):
            raise PermissionDenied
        else :
            return super(AnnounceCreateView, self).dispatch(*args, **kwargs)
    def form_valid(self, form):
        self.object = form.save(commit=False)
        self.object.title = u"[公告]" + self.object.title
        self.object.author_id = self.request.user.id
        self.object.classroom_id = self.kwargs['classroom_id']
        self.object.type = 1
        self.object.save()
        self.object.url = "/teacher/announce/detail/" + str(self.object.id)
        self.object.save()

        #附件
        files = []
        if self.request.FILES.getlist('files'):
             for file in self.request.FILES.getlist('files'):
                fs = FileSystemStorage()
                filename = uuid4().hex
                fs.save("static/attach/"+str(self.request.user.id)+"/"+filename, file)
                files.append([filename, file.name])
        if files:
            for file, name in files:
                content = MessageContent()
                content.title = name
                content.message_id = self.object.id
                content.filename = str(self.request.user.id)+"/"+file
                content.save()
        # 班級學生訊息
        enrolls = Enroll.objects.filter(classroom_id=self.kwargs['classroom_id'])
        for enroll in enrolls:
            messagepoll = MessagePoll(message_id=self.object.id, reader_id=enroll.student_id, message_type=1)
            messagepoll.save()
        return redirect("/teacher/announce/"+self.kwargs['classroom_id'])

    def get_context_data(self, **kwargs):
        context = super(AnnounceCreateView, self).get_context_data(**kwargs)
        context['classroom'] = Classroom.objects.get(id=self.kwargs['classroom_id'])
        return context


# 公告
def announce_detail(request, message_id):
    message = Message.objects.get(id=message_id)
    files = MessageContent.objects.filter(message_id=message_id)
    classroom = Classroom.objects.get(id=message.classroom_id)

    announce_reads = []

    messagepolls = MessagePoll.objects.filter(message_id=message_id)
    for messagepoll in messagepolls:
        try:
            enroll = Enroll.objects.get(classroom_id=message.classroom_id, student_id=messagepoll.reader_id)
            announce_reads.append([enroll.seat, enroll.student.first_name, messagepoll])
        except ObjectDoesNotExist:
            pass

    def getKey(custom):
        return custom[0]
    announce_reads = sorted(announce_reads, key=getKey)
    return render(request, 'teacher/announce_detail.html', {'files':files,'message':message, 'classroom':classroom, 'announce_reads':announce_reads})

# 列出所有課程
class WorkListView(ListView):
    model = Work
    context_object_name = 'works'
    template_name = 'teacher/work_list.html'
    def dispatch(self, *args, **kwargs):
        if not not_in_teacher_group(self.request.user):
            raise PermissionDenied
        else :
            return super(WorkListView, self).dispatch(*args, **kwargs)
    def get_queryset(self):
        if self.kwargs['lesson'] == "1":
            queryset = lesson_list1
        elif self.kwargs['lesson'] == "2":
            queryset = lesson_list2
        elif self.kwargs['lesson'] == "3":
            queryset = lesson_list3
        elif self.kwargs['lesson'] == "4":
            queryset = lesson_list4
        elif self.kwargs['lesson'] == "5":
            queryset = lesson_list2
        elif self.kwargs['lesson'] == "6":
            queryset = lesson_list6       
        elif self.kwargs['lesson'] == "8":
            queryset = lesson_list5  
        elif self.kwargs['lesson'] == "9":
            queryset = lesson_list2   			
        elif self.kwargs['lesson'] == "10":
            queryset = lesson_list7        
        else:
            queryset = lesson_list1
        return queryset

    def get_context_data(self, **kwargs):
        context = super(WorkListView, self).get_context_data(**kwargs)
        context['classroom'] = Classroom.objects.get(id=self.kwargs['classroom_id'])
        context['lesson'] = self.kwargs['lesson']
        return context

# 教師評分
@login_required
@user_passes_test(not_in_teacher_group, login_url='/')
# 列出某作業所有同學名單
def work_class(request, typing, lesson, classroom_id, index):
    enrolls = Enroll.objects.filter(classroom_id=classroom_id)
    classroom = Classroom.objects.get(id=classroom_id)
    classmate_work = []
    scorer_name = ""
    for enroll in enrolls:
        try:
            work = Work.objects.get(typing=0, user_id=enroll.student_id, index=index, lesson_id=lesson)
            if work.scorer > 0 :
                scorer = User.objects.get(id=work.scorer)
                scorer_name = scorer.first_name
            else :
                scorer_name = "1"
        except ObjectDoesNotExist:
            work = Work(typing=0, index=index, user_id=0, lesson_id=lesson)
        except MultipleObjectsReturned:
            work = Work.objects.filter(typing=0, user_id=enroll.student_id, index=index, lesson_id=lesson).last()
        if enroll.group >=0:
            group_name = "第"+str(enroll.group+1)+"組"
        else:
            group_name = "沒有組別"
        assistant = WorkAssistant.objects.filter(typing=0, classroom_id=classroom_id, student_id=enroll.student_id, lesson_id=lesson, index=index)
        if assistant.exists():
            classmate_work.append([enroll,work,1, scorer_name, group_name])
        else :
            classmate_work.append([enroll,work,0, scorer_name, group_name])
    def getKey(custom):
        return custom[0].seat

    classmate_work = sorted(classmate_work, key=getKey)
    return render(request, 'teacher/work_class.html',{'typing':typing, 'classmate_work': classmate_work, 'classroom':classroom, 'index': index, 'lesson':lesson})

# 教師評分
@login_required
@user_passes_test(not_in_teacher_group, login_url='/')
# 列出某作業所有同學名單
def work_group(request, typing, lesson, classroom_id, index):
    enrolls = Enroll.objects.filter(classroom_id=classroom_id)
    classroom = Classroom.objects.get(id=classroom_id)
    classmate_work = []
    scorer_name = ""
    for enroll in enrolls:
        try:
            work = Work.objects.get(typing=0, user_id=enroll.student_id, index=index, lesson_id=lesson)
            if work.scorer > 0 :
                scorer = User.objects.get(id=work.scorer)
                scorer_name = scorer.first_name
            else :
                scorer_name = "1"
        except ObjectDoesNotExist:
            work = Work(typing=0, index=index, user_id=0, lesson_id=lesson)
        except MultipleObjectsReturned:
            work = Work.objects.filter(typing=0, user_id=enroll.student_id, index=index, lesson_id=lesson).last()
        if enroll.group >=0:
            group_name = "第"+str(enroll.group+1)+"組"
        else:
            group_name = "沒有組別"
        assistant = WorkAssistant.objects.filter(typing=0, classroom_id=classroom_id, student_id=enroll.student_id, lesson_id=lesson, index=index)
        if assistant.exists():
            classmate_work.append([enroll,work,1, scorer_name, group_name])
        else :
            classmate_work.append([enroll,work,0, scorer_name, group_name])
    def getKey(custom):
        return custom[0].group, custom[1].publication_date
      
    classmate_work = sorted(classmate_work, key=getKey)    
    return render(request, 'teacher/work_group.html',{'typing':typing, 'classmate_work': classmate_work, 'classroom':classroom, 'index': index, 'lesson':lesson})
  
  
# (小)教師評分
def scoring(request, typing, lesson, classroom_id, user_id, index):
    work_dict = dict(((work.index, [work, WorkFile.objects.filter(work_id=work.id).order_by("-id")]) for work in Work.objects.filter(typing=typing, lesson_id=lesson, user_id=user_id)))
    user = User.objects.get(id=request.user.id)
    teacher = is_teacher(user, classroom_id) or is_assistant(user, classroom_id)
    if not teacher:
        if not WorkAssistant.objects.filter(typing=typing, lesson_id=lesson, index=index, classroom_id=classroom_id, student_id=request.user.id).exists():
            return redirect("/")
    if typing == "0":
        if lesson == "1":
            lesson_name = lesson_list1[int(index)-1][2]
        elif lesson == "2":
            lesson_name = lesson_list2[int(index)-1][1]
        elif lesson == "3":
            lesson_name = lesson_list3[int(index)-1][1]
        elif lesson == "4":
            lesson_name = lesson_list4[int(index)-1][1]
        elif lesson == "5":
            lesson_name = lesson_list2[int(index)-1][1]
        elif lesson == "6":
            lesson_name = lesson_list6[int(index)-1][1]            
        elif lesson == "8":
            lesson_name = lesson_list5[int(index)-1][1]            
        elif lesson == "10":
            lesson_name = lesson_list7[int(index)-1][1]                 
        else:
            lesson_name = lesson_list1[int(index)-1][1]
    elif typing == "1":
        lesson_name = TWork.objects.get(id=index).title
    user = User.objects.get(id=user_id)
    enroll = Enroll.objects.get(classroom_id=classroom_id, student_id=user_id)
    workfiles = []
    try:
        assistant = WorkAssistant.objects.filter(typing=typing, classroom_id=classroom_id,lesson_id=lesson, index=index,student_id=request.user.id)
    except ObjectDoesNotExist:
        if not is_teacher(request.user, classroom_id) or not is_assistant(request.user, classroom_id):
            return render(request, 'message.html', {'message':"您沒有權限"})

    try:
        work3 = Work.objects.get(typing=typing, user_id=user_id, index=index, lesson_id=lesson)
        pic = work3.id
    except ObjectDoesNotExist:
        work3 = Work(typing=typing, index=index, user_id=user_id, lesson_id=lesson)
        pic = 0
    except MultipleObjectsReturned:
        works = Work.objects.filter(typing=typing, user_id=user_id, index=index, lesson_id=lesson).order_by("-id")
        work3 = works[0]
        pic = work3.id
        if int(lesson) > 1 and int(lesson) != 10 :
            prefix = ['static/work/vphysics', 'static/work/euler', 'static/work/ck', 'static/work/vphysics', 'static/work/microbit', 'static/work/pandas', 'static/work/django'][int(lesson) - 2]
            directory = "{prefix}/{uid}/{index}".format(prefix=prefix, uid=user_id, index=index)
            for work in works:
                image_file = "{path}/{id}.jpg".format(path=directory, id=work.id)
                if os.path.exists(image_file):
                    pic = work.id
                    break

    if request.method == 'POST':
        if lesson == "4":
            form = ScoreBForm(request.user, request.POST)
        else :
            form = ScoreForm(request.user, request.POST)
        if form.is_valid():
            works = Work.objects.filter(typing=typing, index=index, user_id=user_id, lesson_id=lesson)
            if works.exists():
                if works[0].score < 0 :
                        # 小老師
                        if not is_teacher(request.user, classroom_id):
                            # credit
                            update_avatar(request.user.id, 2, 1)
                            # History
                            history = PointHistory(user_id=request.user.id, kind=2, message=u'1分--小老師:<'+lesson_name.decode('utf8')+u'><'+enroll.student.first_name+u'>', url="/student/work/show/"+lesson+"/"+index)
                            history.save()

                        # credit
                        update_avatar(enroll.student_id, 1, 1)
                        # History
                        history = PointHistory(user_id=user_id, kind=1, message=u'1分--作業受評<'+lesson_name.decode('utf8')+u'><'+request.user.first_name+u'>', url="/student/work/show/"+lesson+"/"+index)
                        history.save()

                works.update(score=form.cleaned_data['score'])
                works.update(scorer=request.user.id)
                works.update(comment=form.cleaned_data['comment'])
				
                if form.cleaned_data['comment']:
                    # create Message
                    title = u"<" + request.user.first_name+ u">給了評語<" + lesson_name.decode('utf8') + u">"
                    url = "/student/work/show/" + typing + "/" + lesson + "/" + index + "/" + str(enroll.student_id)
                    message = Message(title=title, url=url, time=timezone.now())
                    message.save()		

                    # message for group member
                    messagepoll = MessagePoll(message_id = message.id,reader_id=enroll.student_id)
                    messagepoll.save()					

            if is_teacher(request.user, classroom_id) or is_assistant(request.user, classroom_id):
                if form.cleaned_data['assistant']:
                    try :
                        assistant = WorkAssistant.objects.get(typing=typing, student_id=user_id, classroom_id=classroom_id, index=index, lesson_id=lesson)
                    except ObjectDoesNotExist:
                        assistant = WorkAssistant(typing=typing, student_id=user_id, classroom_id=classroom_id, index=index, lesson_id=lesson)
                        assistant.save()

                    # create Message
                    title = u"<" + assistant.student.first_name+ u">擔任小老師<" + lesson_name.decode('utf8') + u">"
                    url = "/teacher/score_peer/" + typing + "/" + lesson + "/" + index + "/" + classroom_id + "/" + str(enroll.group)
                    message = Message(title=title, url=url, time=timezone.now())
                    message.save()

                    group = Enroll.objects.get(classroom_id=classroom_id, student_id=assistant.student_id).group
                    if group >= 0 :
                        enrolls = Enroll.objects.filter(group = group)
                        for enroll in enrolls:
                            # message for group member
                            messagepoll = MessagePoll(message_id = message.id,reader_id=enroll.student_id)
                            messagepoll.save()
                if typing == "0":
                    return redirect('/teacher/work/class/'+typing+ "/" + lesson+"/"+classroom_id+'/'+index)
                elif typing == "1":
                    return redirect('/teacher/work2/class/' + lesson+"/"+classroom_id+'/'+index)  
            else:
                return redirect('/teacher/score_peer/'+typing+"/"+lesson+"/"+index+'/'+classroom_id+'/'+str(enroll.group))

    else:
        works = Work.objects.filter(typing=typing, index=index, user_id=user_id, lesson_id=lesson).order_by("-id")
        if not works.exists():
            if lesson == "4":
               form = ScoreBForm(user=request.user)
            else :
               form = ScoreForm(user=request.user)
        else:
            if lesson == "4":
                form = ScoreBForm(instance=works[0], user=request.user)
            else :
                form = ScoreForm(instance=works[0], user=request.user)
            workfiles = WorkFile.objects.filter(work_id=works[0].id).order_by("-id")
    return render(request, 'teacher/scoring.html', {'typing':typing, 'form': form,'work':work3, 'pic':pic, 'workfiles':workfiles, 'teacher':teacher, 'student':user, 'classroom_id':classroom_id, 'lesson':lesson, 'index':index, 'work_dict':work_dict})

# 小老師評分名單
def score_peer(request, typing, lesson, index, classroom_id, group):
    if typing == "0":
        if lesson == "1":
            queryset = lesson_list1
        elif lesson == "2":
            queryset = lesson_list2
        elif lesson == "3":
            queryset = lesson_list3
        elif lesson == "4":
            queryset = lesson_list4
        elif lesson == "5":
            queryset = lesson_list2 
        elif lesson == "6":
            queryset = lesson_list6                     
        else:
            queryset = lesson_list1
    elif typing == "1":
        queryset = TWork.objects.filter(classroom_id=classroom_id)
    try:
        assistant = WorkAssistant.objects.get(typing=typing, lesson_id=lesson, index=index, classroom_id=classroom_id, student_id=request.user.id)
    except ObjectDoesNotExist:
        if typing == "0":
            return redirect("/student/work/group/0/"+lesson+"/"+index+"/"+classroom_id+"/#"+index)
        elif typing == "1":
            return redirect("/student/work/group/1/"+lesson+"/"+index+"/"+classroom_id+"/#"+index)
    enrolls = Enroll.objects.filter(classroom_id=classroom_id, group=group)
    lessons = ""
    classmate_work = []
    for enroll in enrolls:
        if not enroll.student_id == request.user.id :
            scorer_name = ""
            try:
                work = Work.objects.get(typing=typing, user_id=enroll.student.id, index=index, lesson_id=lesson)
                if work.scorer > 0 :
                    scorer = User.objects.get(id=work.scorer)
                    scorer_name = scorer.first_name
            except ObjectDoesNotExist:
                work = Work(typing=typing, index=index, user_id=enroll.student.id, lesson_id=lesson)
            except MultipleObjectsReturned:
                work = Work.objects.filter(typing=typing, user_id=enroll.student.id, index=index, lesson_id=lesson).order_by("-id")[0]
            classmate_work.append([enroll.student,work,1, scorer_name])
        lessons = queryset[int(index)-1]
    return render(request, 'teacher/score_peer.html',{'enrolls':enrolls, 'classmate_work': classmate_work, 'classroom_id':classroom_id, 'lesson':lesson, 'index': index, 'typing':typing})

# 心得
def memo(request, lesson, classroom_id):
    # 限本班任課教師
    if not is_teacher(request.user, classroom_id) and not is_assistant(request.user, classroom_id):
        return redirect("/")
    enrolls = Enroll.objects.filter(classroom_id=classroom_id).order_by("seat")
    classroom_name = Classroom.objects.get(id=classroom_id).name
    return render(request, 'teacher/memo.html', {'lesson':lesson, 'enrolls':enrolls, 'classroom_name':classroom_name})

# 評分某同學某進度心得
@login_required
@user_passes_test(not_in_teacher_group, login_url='/')
def check(request, typing, lesson, unit, user_id, classroom_id):
    # 限本班任課教師
    if not is_teacher(request.user, classroom_id) and not is_assistant(request.user, classroom_id):
        return redirect("/")

    user_name = User.objects.get(id=user_id).first_name
    lesson_dict = {}
    works = Work.objects.filter(typing=typing, user_id=user_id, lesson_id=lesson)
    enroll = Enroll.objects.get(student_id=user_id, classroom_id=classroom_id)
    if typing == "0":
        if lesson == "1":
            for index,assignment in enumerate(lesson_list1):
                lesson_dict[index] = [assignment]
        elif lesson == "2" :
            for index,assignment in enumerate(lesson_list2):
                lesson_dict[index] = [assignment]
        elif lesson == "3" :
            for index,assignment in enumerate(lesson_list3):
                lesson_dict[index] = [assignment]
        elif lesson == "4" :
            for index,assignment in enumerate(lesson_list4):
                lesson_dict[index] = [assignment]
        elif lesson == "5" :
            for index,assignment in enumerate(lesson_list2):
                lesson_dict[index] = [assignment]
        elif lesson == "6" :
            for index,assignment in enumerate(lesson_list6):
                lesson_dict[index] = [assignment]                
        elif lesson == "8" :
            for index,assignment in enumerate(lesson_list5):
                lesson_dict[index] = [assignment]                
    else :
        assignments = TWork.objects.filter(classroom_id=classroom_id)
        for assignment in assignments:
            lesson_dict[assignment.id] = [assignment]

    for work in works:
        if typing == "0":
            index = work.index - 1
        else:
            index = work.index
        if index in lesson_dict:
            lesson_dict[index].append(work.score)
            lesson_dict[index].append(work.publication_date)
            if work.score > 0:
                score_name = User.objects.get(id=work.scorer).first_name
                lesson_dict[index].append(score_name)
            else :
                lesson_dict[index].append("尚未評分!")
            lesson_dict[index].append(work.memo)
    if request.method == 'POST':
      if typing == "0":
        if lesson == "1":
            if unit == "1":
                form = CheckForm1(request.POST)
            elif unit == "2":
                form = CheckForm2(request.POST)
            elif unit == "3":
                form = CheckForm3(request.POST)
            elif unit == "4":
                form = CheckForm4(request.POST)
        elif lesson == "2":
            form = CheckForm_vphysics(request.POST)
        elif lesson == "3":
            form = CheckForm_euler(request.POST)
        elif lesson == "4":
            form = CheckForm_vphysics(request.POST)
        elif lesson == "5":
            form = CheckForm_vphysics(request.POST)
        elif lesson == "6":
            form = CheckForm_microbit(request.POST)
        elif lesson == "7":
            form = CheckForm_pandas(request.POST)   
        elif lesson == "8":
            form = CheckForm_django(request.POST)                 
      elif typing == "1":
          form = CheckFormCustom(request.POST)                               
      else:
          form = CheckForm(request.POST)

      if form.is_valid():
        if typing == "0":
          if lesson == "1":
              if unit == "1":
                  enroll.score_memo1=form.cleaned_data['score_memo1']
              elif unit == "2":
                  enroll.score_memo2=form.cleaned_data['score_memo2']
              elif unit == "3":
                  enroll.score_memo3=form.cleaned_data['score_memo3']
              elif unit == "4":
                  enroll.score_memo4=form.cleaned_data['score_memo4']    
          elif lesson == "2" :
              enroll.score_memo_vphysics=form.cleaned_data['score_memo_vphysics']
          elif lesson == "3":
              enroll.score_memo_euler=form.cleaned_data['score_memo_euler']
          elif lesson == "4" :
              enroll.score_memo_vphysics2=form.cleaned_data['score_memo_vphysics']
          elif lesson == "5" :
              enroll.score_memo_vphysics3=form.cleaned_data['score_memo_vphysics']
          elif lesson == "6" :
              enroll.score_memo_microbit=form.cleaned_data['score_memo_microbit']
          elif lesson == "7" :
              enroll.score_memo_pandas=form.cleaned_data['score_memo_pandas']                            
          elif lesson == "8" :
              enroll.score_memo_django=form.cleaned_data['score_memo_django']                
          enroll.save()
          #if form.cleaned_data['certificate']:
          #    return redirect('/certificate/'+lesson+'/'+unit+'/'+str(enroll.id)+'/certificate')
          #else:
          return redirect('/teacher/memo/'+lesson+"/"+classroom_id)
        elif typing == "1":
          enroll.score_memo_custom = form.cleaned_data['score_memo_custom']
          enroll.save()    
          return redirect('/teacher/memo/'+lesson+"/"+classroom_id)                 
        else :
          enroll.score_memo = form.cleaned_data['score_memo']
          enroll.save()
          return redirect('/teacher/memo/'+lesson+"/"+classroom_id)
    else:
      if typing == "0":
        if lesson == "1":
            if unit == "1":
                form = CheckForm1(instance=enroll)
            elif unit == "2":
                form = CheckForm2(instance=enroll)
            elif unit == "3":
                form = CheckForm3(instance=enroll)
            elif unit == "4":
                form = CheckForm4(instance=enroll)
        elif lesson == "2" or lesson == "4" or lesson == "5":
            form = CheckForm_vphysics(instance=enroll)
        elif lesson == "3":
            form = CheckForm_euler(instance=enroll)
        elif lesson == "6":
            form = CheckForm_microbit(instance=enroll)
        elif lesson == "7":
            form = CheckForm_pandas(instance=enroll)                        
        elif lesson == "8":
            form = CheckForm_django(instance=enroll)            
        else :
            form =  CheckForm1(instance=enroll)
      elif typing == "1":
        form = CheckFormCustom(instance=enroll)
      else:
        form =  CheckForm(instance=enroll)

    return render(request, 'teacher/check.html', {'typing':typing, 'works':works, 'lesson': lesson, 'unit':unit, 'form':form, 'works':works, 'lesson_list':sorted(lesson_dict.items()), 'enroll': enroll, 'classroom_id':classroom_id})

@login_required
@user_passes_test(not_in_teacher_group, login_url='/')
def grade(request, typing, lesson, unit, classroom_id):
    # 限本班任課教師
    if not is_teacher(request.user, classroom_id) and not is_assistant(request.user, classroom_id):
        return redirect("/")
    enrolls = Enroll.objects.filter(classroom_id=classroom_id).order_by('seat')
    classroom = Classroom.objects.get(id=classroom_id)     
    user_ids = [enroll.student_id for enroll in enrolls]
    work_pool = Work.objects.filter(typing=typing, user_id__in=user_ids, lesson_id=lesson).order_by('id')
    lesson_dict = {}
    data = []
    lesson_list = [lesson_list1, lesson_list2, lesson_list3, lesson_list4, lesson_list2, lesson_list6, lesson_list2, lesson_list5][int(lesson)-1]
    for enroll in enrolls:
      enroll_score = []
      enroll_grade = []
      total = 0
      memo = 0 
      grade = 0
      stu_works = filter(lambda w: w.user_id == enroll.student_id, work_pool)
      if typing == "0":
        if lesson == "1":
            if unit == "1":
                lesson_list = lesson_list[0:17]			
            elif unit == "2":
                lesson_list = lesson_list[17:25]
            elif unit == "3":
                lesson_list = lesson_list[25:33]
            elif unit == "4":
                lesson_list = lesson_list[33:41]
      elif typing == "1":
        lesson_list = TWork.objects.filter(classroom_id=classroom_id)
      else :
        lesson_list = CWork.objects.filter(classroom_id=classroom_id)
      for index, assignment in enumerate(lesson_list):
            if typing == "0": 			    
                if unit == "1":
                    work_index = index + 1
                elif unit == "2":
                    work_index = index + 1 + 17
                elif unit == "3":
                    work_index = index + 1 + 17 + 8
                elif unit == "4":
                    work_index = index + 1 + 17 + 8 + 8
                works = list(filter(lambda w: w.index == work_index, stu_works))
            else :
                works = list(filter(lambda w: w.index == assignment.id, stu_works))
                work_index = index + 1
            works_count = len(works)
            if works_count == 0:
                enroll_score.append(["X", work_index])
                if typing == "0" or typing == "1":
                    if not lesson == "4":
                        total += 60
            else:
                work = works[-1]
                enroll_score.append([work.score, index])
                if work.score == -2:
                    if typing == "0" or typing == "1":
                          if not lesson == "4":
                              total += 80
                else:
                    total += work.score

            if typing == "0":
                if lesson == "1":
                    memo = [enroll.score_memo1, enroll.score_memo2, enroll.score_memo3, enroll.score_memo4][int(unit)-1]
                elif lesson == "2":
                    memo = enroll.score_memo_vphysics
                elif lesson == "3":
                    memo = enroll.score_memo_euler
                elif lesson == "4":
                    memo = enroll.score_memo_vphysics2
                elif lesson == "5":
                    memo = enroll.score_memo_vphysics3
                elif lesson == "6":
                    memo = enroll.score_memo_microbit
                elif lesson == "7":
                    memo = enroll.score_memo_pandas
                elif lesson == "8":
                    memo = enroll.score_memo_django                                                            
                grade = int(total / len(lesson_list) * 0.6 + memo * 0.4)					
            elif typing == "1":
                memo = enroll.score_memo_custom
                grade = int(total / len(lesson_list) * 0.6 + memo * 0.4)
            elif typing == "2":
                memo = enroll.score_memo_check
                grade = total
      data.append([enroll, enroll_score, memo, grade])
    return render(request, 'teacher/grade.html', {'typing':typing, 'lesson':lesson, 'unit':unit, 'lesson_list':lesson_list, 'classroom':classroom, 'data':data})

@login_required
@user_passes_test(not_in_teacher_group, login_url='/')
def grade_excel(request, typing, lesson, unit, classroom_id):
    # 限本班任課教師
    if not is_teacher(request.user, classroom_id) and not is_assistant(request.user, classroom_id):
        return redirect("/")
    enrolls = Enroll.objects.filter(classroom_id=classroom_id).order_by('seat')
    user_ids = [enroll.student_id for enroll in enrolls]
    work_pool = Work.objects.filter(typing=typing, user_id__in=user_ids, lesson_id=lesson).order_by('id')
    lesson_dict = {}
    data = []
    lesson_list = [lesson_list1, lesson_list2, lesson_list3, lesson_list4, lesson_list2, lesson_list6, lesson_list2, lesson_list5][int(lesson)-1]
    for enroll in enrolls:
      enroll_score = []
      total = 0
      stu_works = filter(lambda w: w.user_id == enroll.student_id, work_pool)
      if typing == "0":
        if lesson == "1":
            if unit == "1":
                lesson_list = lesson_list[0:17]
            elif unit == "2":
                lesson_list = lesson_list[17:25]
            elif unit == "3":
                lesson_list = lesson_list[25:33]
            elif unit == "4":
                lesson_list = lesson_list[33:41]
      elif typing == "1":
        lesson_list = TWork.objects.filter(classroom_id=classroom_id)
      else :
        lesson_list = CWork.objects.filter(classroom_id=classroom_id)
      memo = ""
      grade = 0
      for index, assignment in enumerate(lesson_list):
            if typing == "0": 
                works = filter(lambda w: w.index == index+1, stu_works)
            else :
                works = filter(lambda w: w.index == assignment.id, stu_works)
            works_count = len(works)
            if works_count == 0:
                enroll_score.append(["X", index])
                if typing == "0" or typing == "1":
                    if not lesson == "4":
                        total += 60
            else:
                work = works[-1]
                enroll_score.append([work.score, index])
                if work.score == -2:
                    if typing == "0" or typing == "1":
                          if not lesson == "4":
                              total += 80
                else:
                    total += work.score

            if typing == "0":
                if lesson == "1":
                    memo = [enroll.score_memo1, enroll.score_memo2, enroll.score_memo3, enroll.score_memo4][int(unit)-1]
                elif lesson == "2":
                    memo = enroll.score_memo_vphysics
                elif lesson == "3":
                    memo = enroll.score_memo_euler
                elif lesson == "4":
                    memo = enroll.score_memo_vphysics2
                elif lesson == "5":
                    memo = enroll.score_memo_vphysics3
                elif lesson == "6":
                    memo = enroll.score_memo_microbit 
                elif lesson == "7":
                    memo = enroll.score_memo_pandas                         
                elif lesson == "8":
                    memo = enroll.score_memo_django
            elif typing == "1":
                memo = enroll.score_memo_custom
            if typing == "2":
                grade = total
            else :
                grade = int(total / len(lesson_list) * 0.6 + memo * 0.4)
      data.append([enroll, enroll_score, memo, grade])
                
    classroom = Classroom.objects.get(id=classroom_id)       
    output = StringIO.StringIO()
    workbook = xlsxwriter.Workbook(output)    
    worksheet = workbook.add_worksheet(classroom.name)
    date_format = workbook.add_format({'num_format': 'yy/mm/dd'})

    row = 1
    worksheet.write(row, 1, u'座號')
    worksheet.write(row, 2, u'姓名')
    worksheet.write(row, 3, u'成績')        
    worksheet.write(row, 4, u'心得')         
    index = 5
    for assignment in lesson_list:
        if typing == "0":
	          worksheet.write(row, index, assignment[1])
        else :
	          worksheet.write(row, index, assignment.title)              
        index += 1

    index = 5
    if not typing == "0":
        row += 1
        for assignment in lesson_list:            
            worksheet.write(row, index, datetime.strptime(str(assignment.time)[:19],'%Y-%m-%d %H:%M:%S'), date_format)
            index += 1			

    for enroll, enroll_score, memo, grade in data:
      row += 1
      worksheet.write(row, 1, enroll.seat)
      worksheet.write(row, 2, enroll.student.first_name)
      worksheet.write(row, 3, grade)         
      worksheet.write(row, 4, memo)     
      index = 5
      for score, index2 in enroll_score:
          if score == -2 :
              worksheet.write(row, index, "V")
          else :
              worksheet.write(row, index, score)
          index +=1 

    workbook.close()
    # xlsx_data contains the Excel file
    response = HttpResponse(content_type='application/vnd.ms-excel')
    if typing == "0":
        type_name = "指定作業"
    elif typing == "1":
        type_name = "自訂作業"
    else:
        type_name = "檢核作業"        
    filename = classroom.name + '-' + type_name + "-" + str(localtime(timezone.now()).date()) + '.xlsx'
    response['Content-Disposition'] = 'attachment; filename={0}'.format(filename.encode('utf8'))
    xlsx_data = output.getvalue()
    response.write(xlsx_data)
    return response
  
# 列出分組所有作業
@login_required
@user_passes_test(not_in_teacher_group, login_url='/')
def work1(request, lesson, classroom_id):
    classroom_name = Classroom.objects.get(id=classroom_id).name
    lessons = []
    lesson_dict = {}
    groups = [group for group in EnrollGroup.objects.filter(classroom_id=classroom_id)]
    enroll_pool = [enroll for enroll in Enroll.objects.filter(classroom_id=classroom_id).order_by('seat')]
    student_ids = map(lambda a: a.student_id, enroll_pool)
    work_pool = Work.objects.filter(user_id__in=student_ids, lesson_id=lesson)
    user_pool = [user for user in User.objects.filter(id__in=work_pool.values('scorer'))]
    assistant_pool = [assistant for assistant in Assistant.objects.filter(classroom_id=classroom_id)]
    if lesson == "1":
        assignments = lesson_list
    else :
        assignments = TWork.objects.filter(classroom_id=classroom_id)
    for index,assignment in enumerate(assignments):
        works = []
        if lesson == "1":
            lesson_dict[index] = [assignment[1]]
        else:
            lesson_dict[index] = [assignment.title]
        student_groups = []
        for group in groups:
            members = filter(lambda u: u.group == group.id, enroll_pool)
            group_assistants = []
            scorer_name = ""
            for member in members:
                if lesson == "1":
                    sworks = filter(lambda w: w.index == index and w.student_id == member.student_id, work_pool)
                else:
                    sworks = filter(lambda w: w.index == assignment.id and w.student_id == enroll.student_id, work_pool)
                if sworks:
                    work = sworks[-1]
                    scorer = filter(lambda u: u.id == work.scorer, user_pool)
                    scorer_name = scorer[0].first_name if scorer else 'X'
                else:
                    work = SWork(index=index, student_id=1)
                works.append([member, work.score, scorer_name, work.memo])
                assistant = filter(lambda a: a.student_id == member.student_id and a.index == index, assistant_pool)
                if assistant:
                    group_assistants.append(member)
            student_groups.append([group, works, group_assistants])
        lessons.append([lesson_list[index], student_groups])
    return render(request, 'teacher/work1.html', {'lesson':lesson, 'lessons':lessons, 'classroom_id':classroom_id})

# Ajax 設為小教師、取消小教師
def make_work_assistant(request):
    classroom_id = request.POST.get('classroomid')
    user_id = request.POST.get('userid')
    action = request.POST.get('action')
    lesson = request.POST.get('lesson')
    index = request.POST.get('index')
    if lesson == "1":
        queryset = lesson_list1
        assignment = queryset[int(index)-1][2]
    elif lesson == "2":
        queryset = lesson_list2
        assignment = queryset[int(index)-1][1]
    elif lesson == "3":
        queryset = lesson_list3
        assignment = queryset[int(index)-1][1]
    else:
        queryset = lesson_list1
        assignment = queryset[int(index)-1][2]

    if is_teacher(request.user, classroom_id) or is_assistant(request.user, classroom_id):
        if user_id and action and lesson and index:
            user = User.objects.get(id=user_id)
        if action == "set":
            try:
                assistant = WorkAssistant.objects.get(student_id=user_id, classroom_id=classroom_id, lesson_id=lesson, index=index)
            except ObjectDoesNotExist:
                assistant = WorkAssistant(student_id=user_id, classroom_id=classroom_id, lesson_id=lesson, index=index)
            assistant.save()

            # create Message
            group = Enroll.objects.get(classroom_id=classroom_id, student_id=assistant.student_id).group
            title = "<" + assistant.student.first_name.encode("utf-8") + u">擔任小老師<".encode("utf-8") + assignment + ">"
            url = "/teacher/score_peer/" + str(lesson) + "/" + index + "/"+ classroom_id + "/" + str(group)

            message = Message(title=title, url=url, time=timezone.now())
            message.save()
        else:
            try:
                assistant = WorkAssistant.objects.get(student_id=user_id, classroom_id=classroom_id, lesson_id=lesson, index=index)
                assistant.delete()
            except ObjectDoesNotExist:
                pass
            # create Message
            title = "<" + assistant.student.first_name.encode("utf-8") + u">取消小老師<".encode("utf-8") + assignment + ">"
            url = "/student/group/work/" + lesson + "/" + index + "/" + classroom_id
            message = Message(title=title, url=url, time=timezone.now())
            message.save()

        group = Enroll.objects.get(classroom_id=classroom_id, student_id=assistant.student_id).group
        if group > 0 :
            enrolls = Enroll.objects.filter(group = group)
            for enroll in enrolls:
                # message for group member
                messagepoll = MessagePoll(message_id = message.id,reader_id=enroll.student_id)
                messagepoll.save()
        return JsonResponse({'status':'ok'}, safe=False)
    else:
        return JsonResponse({'status':classroom_id}, safe=False)
# Create your views here.
def import_sheet(request):
    if False:
        return redirect("/")
    if request.method == "POST":
        form = UploadFileForm(request.POST,
                              request.FILES)
        if form.is_valid():
            ImportUser.objects.all().delete()
            request.FILES['file'].save_to_database(
                name_columns_by_row=0,
                model=ImportUser,
                mapdict=['username', 'first_name', 'password'])
            users = ImportUser.objects.all()
            return render(request, 'teacher/import_student.html',{'users':users})
        else:
            return HttpResponseBadRequest()
    else:
        form = UploadFileForm()
    return render(
        request,
        'teacher/import_form.html',
        {
            'form': form,
            'title': 'Excel file upload and download example',
            'header': ('Please choose any excel file ' +
                       'from your cloned repository:')
        })

# Create your views here.
def import_student(request):
    if False:
        return redirect("/")

    users = ImportUser.objects.all()
    username_list = [request.user.username+"_"+user.username for user in users]
    exist_users = [user.username for user in User.objects.filter(username__in=username_list)]
    create_list = []
    for user in users:
        username = request.user.username+"_"+user.username
        if username in exist_users:
            continue
        new_user = User(username=username, first_name=user.first_name, last_name=request.user.last_name, password=user.password, email=username+"@edu.tw")
        new_user.set_password(user.password)
        create_list.append(new_user)

    User.objects.bulk_create(create_list)
    new_users = User.objects.filter(username__in=[user.username for user in create_list])

    profile_list = []
    message_list = []
    poll_list = []
    title = "請洽詢任課教師課程名稱及選課密碼"
    url = "/student/classroom/add"
    message = Message(title=title, url=url, time=timezone.now())
    message.save()
    for user in new_users:
        profile = Profile(user=user)
        profile_list.append(profile)
        poll = MessagePoll(message_id=message.id, reader_id=user.id)
        poll_list.append(poll)

    Profile.objects.bulk_create(profile_list)
    MessagePoll.objects.bulk_create(poll_list)

    return redirect('/teacher/student/list')

# 教師可以查看所有帳號
class StudentListView(ListView):
    context_object_name = 'users'
    paginate_by = 50
    template_name = 'teacher/student_list.html'

    def get_queryset(self):
        username = username__icontains=self.request.user.username+"_"
        if self.request.GET.get('account') != None:
            keyword = self.request.GET.get('account')
            queryset = User.objects.filter(Q(username__icontains=username+keyword) | (Q(first_name__icontains=keyword) & Q(username__icontains=username))).order_by('-id')
        else :
            queryset = User.objects.filter(username__icontains=username).order_by('-id')
        return queryset

    def get_context_data(self, **kwargs):
        context = super(StudentListView, self).get_context_data(**kwargs)
        account = self.request.GET.get('account')
        context.update({'account': account})
        return context

# 修改密碼
def password(request, user_id):
    user = User.objects.get(id=user_id)
    if not user.username.startswith(request.user.username):
        return redirect("/")
    if request.method == 'POST':
        form = PasswordForm(request.POST)
        if form.is_valid():
            user = User.objects.get(id=user_id)
            user.set_password(request.POST['password'])
            user.save()
            return redirect('/teacher/student/list/')
    else:
        form = PasswordForm()
        user = User.objects.get(id=user_id)

    return render(request, 'form.html',{'form': form, 'user':user})

# 修改真實姓名
def realname(request, user_id):
    user = User.objects.get(id=user_id)
    if not user.username.startswith(request.user.username):
        return redirect("/")
    if request.method == 'POST':
        form = RealnameForm(request.POST)
        if form.is_valid():
            user = User.objects.get(id=user_id)
            user.first_name =form.cleaned_data['first_name']
            user.save()
            return redirect('/teacher/student/list/')
    else:
        teacher = False
        enrolls = Enroll.objects.filter(student_id=user_id)
        for enroll in enrolls:
            classroom = Classroom.objects.get(id=enroll.classroom_id)
            if request.user.id == classroom.teacher_id:
                teacher = True
                break
        if teacher or request.user.is_superuser:
            user = User.objects.get(id=user_id)
            form = RealnameForm(instance=user)
        else:
            return redirect("/")

    return render(request, 'form.html',{'form': form})

# 列出所有課程
class WorkListView2(ListView):
    model = TWork
    context_object_name = 'works'
    template_name = 'teacher/twork_list.html'
    paginate_by = 20

    def get_queryset(self):
        queryset = TWork.objects.filter(classroom_id=self.kwargs['classroom_id']).order_by("-id")
        return queryset

    def get_context_data(self, **kwargs):
        context = super(WorkListView2, self).get_context_data(**kwargs)
        context['classroom'] = Classroom.objects.get(id=self.kwargs['classroom_id'])
        context['lesson'] = self.kwargs['lesson']
        return context

#新增一個課程
class WorkCreateView2(CreateView):
    model = TWork
    form_class = WorkForm
    template_name = 'form.html'
    def form_valid(self, form):
        self.object = form.save(commit=False)
        self.object.teacher_id = self.request.user.id
        self.object.classroom_id = self.kwargs['classroom_id']
        self.object.save()
        return redirect("/teacher/work2/"+self.kwargs['lesson']+"/"+self.kwargs['classroom_id'])

# 修改選課密碼
def work_edit(request, classroom_id):
    # 限本班任課教師
    if not is_teacher(request.user, classroom_id):
        return redirect("/")
    classroom = Classroom.objects.get(id=classroom_id)
    if request.method == 'POST':
        form = ClassroomForm(request.POST)
        if form.is_valid():
            classroom.name =form.cleaned_data['name']
            classroom.password = form.cleaned_data['password']
            classroom.save()
            # 記錄系統事件
            if is_event_open(request) :
                log = Log(user_id=request.user.id, event=u'修改選課密碼<'+classroom.name+'>')
                log.save()
            return redirect('/teacher/classroom')
    else:
        form = ClassroomForm(instance=classroom)

    return render(request, 'form.html',{'form': form})

# 列出某作業所有同學名單
def work_class2(request, lesson, classroom_id, work_id):
    enrolls = Enroll.objects.filter(classroom_id=classroom_id)
    classroom = Classroom.objects.get(id=classroom_id)
    classmate_work = []
    scorer_name = ""
    for enroll in enrolls:
        try:
            work = Work.objects.get(typing=1, user_id=enroll.student_id, index=work_id, lesson_id=lesson)
            if work.scorer > 0 :
                scorer = User.objects.get(id=work.scorer)
                scorer_name = scorer.first_name
            else :
                scorer_name = "1"
        except ObjectDoesNotExist:
            work = Work(typing=1, index=work_id, user_id=0, lesson_id=lesson)
        except MultipleObjectsReturned:
            work = Work.objects.filter(typing=1, user_id=enroll.student_id, index=work_id, lesson_id=lesson).last()
        assistant = WorkAssistant.objects.filter(typing=1, classroom_id=classroom_id, student_id=enroll.student_id, lesson_id=lesson, index=work_id)
        if enroll.group >=0:
            group_name = "第"+str(enroll.group+1)+"組"
        else:
            group_name = "沒有組別"
        if assistant.exists():
            classmate_work.append([enroll,work,1, scorer_name, group_name])
        else :
            classmate_work.append([enroll,work,0, scorer_name, group_name])
    def getKey(custom):
        return custom[0].seat

    classmate_work = sorted(classmate_work, key=getKey)

    return render(request, 'teacher/work_class.html',{'typing':1, 'classmate_work': classmate_work, 'classroom':classroom, 'index': work_id})

# 設定班級助教
def classroom_assistant(request, classroom_id):
    # 限本班任課教師
    if not is_teacher(request.user, classroom_id):
        return redirect("/")
    assistants = Assistant.objects.filter(classroom_id=classroom_id).order_by("-id")
    classroom = Classroom.objects.get(id=classroom_id)

    return render(request, 'teacher/assistant.html',{'assistants': assistants, 'classroom':classroom})

# 教師可以查看所有帳號
class AssistantListView(ListView):
    context_object_name = 'users'
    paginate_by = 20
    template_name = 'teacher/assistant_user.html'

    def get_queryset(self):
        if self.request.GET.get('account') != None:
            keyword = self.request.GET.get('account')
            queryset = User.objects.filter(Q(groups__name__in=['apply']) & (Q(username__icontains=keyword) | Q(first_name__icontains=keyword))).order_by('-id')
        else :
            queryset = User.objects.filter(groups__name__in=['apply']).order_by('-id')
        return queryset

    def get_context_data(self, **kwargs):
        context = super(AssistantListView, self).get_context_data(**kwargs)
        context['classroom'] = Classroom.objects.get(id=self.kwargs['classroom_id'])
        assistant_list = []
        assistants = Assistant.objects.filter(classroom_id=self.kwargs['classroom_id'])
        for assistant in assistants:
            assistant_list.append(assistant.user_id)
        context['assistants'] = assistant_list
        return context

# 列出所有助教課程
class AssistantClassroomListView(ListView):
    model = Classroom
    context_object_name = 'classrooms'
    template_name = 'teacher/assistant_list.html'
    paginate_by = 20
    def get_queryset(self):
        assistants = Assistant.objects.filter(user_id=self.request.user.id)
        classroom_list = []
        for assistant in assistants:
            classroom_list.append(assistant.classroom_id)
        queryset = Classroom.objects.filter(id__in=classroom_list).order_by("-id")
        return queryset

# 列出所有課程
class WorkListView3(ListView):
    model = CWork
    context_object_name = 'works'
    template_name = 'teacher/cwork_list.html'
    paginate_by = 20

    def dispatch(self, *args, **kwargs):
        if not not_in_teacher_group(self.request.user):
            raise PermissionDenied
        else :
            return super(WorkListView3, self).dispatch(*args, **kwargs)

    def get_queryset(self):
        queryset = CWork.objects.filter(classroom_id=self.kwargs['classroom_id']).order_by("-id")
        return queryset

    def get_context_data(self, **kwargs):
        context = super(WorkListView3, self).get_context_data(**kwargs)
        context['classroom'] = Classroom.objects.get(id=self.kwargs['classroom_id'])
        context['lesson'] = self.kwargs['lesson']
        return context

#新增一個課程
class WorkCreateView3(CreateView):
    model = CWork
    form_class = Work3Form
    template_name = 'form.html'

    def dispatch(self, *args, **kwargs):
        if not not_in_teacher_group(self.request.user):
            raise PermissionDenied
        else :
            return super(WorkCreateView3, self).dispatch(*args, **kwargs)
    def form_valid(self, form):
        self.object = form.save(commit=False)
        self.object.teacher_id = self.request.user.id
        self.object.classroom_id = self.kwargs['classroom_id']
        self.object.save()
        return redirect("/teacher/work3/"+self.kwargs['lesson']+"/"+self.kwargs['classroom_id'])


# 列出某作業所有同學名單
def work_class3(request, lesson, classroom_id, work_id):
    if not not_in_teacher_group(request.user):
        return redirect("/")
    enrolls = Enroll.objects.filter(classroom_id=classroom_id).order_by('seat')
    classroom = Classroom.objects.get(id=classroom_id)
    classmate_work = []
    groups = {}
    scorer_name = ""
    for enroll in enrolls:
        try:
            work = Work.objects.get(typing=2, user_id=enroll.student_id, index=work_id, lesson_id=lesson)
        except ObjectDoesNotExist:
            work = Work(typing=2, user_id=enroll.student_id, index=work_id, lesson_id=lesson, score=0)
        except MultipleObjectsReturned:
            work = Work.objects.filter(typing=2, user_id=enroll.student_id, index=work_id, lesson_id=lesson).last()
        if enroll.group >=0:
            group_name = "第"+str(enroll.group+1)+"組"
        else:
            group_name = "沒有組別"
        classmate_work.append([enroll,work, group_name])
    return render(request, 'teacher/work3_class.html',{'typing':2, 'classmate_work': classmate_work, 'classroom':classroom, 'index': work_id, 'groups': groups})

#
def work3_score(request, lesson, classroom_id, work_id):
    # 檢查是否為教師或助教
    assistants = Assistant.objects.filter(classroom_id=classroom_id, user_id=request.user.id)
    if not is_teacher(request.user, classroom_id): # and not assistants:
        return JsonResponse({'status':'fail'}, safe=False)
    
    score = request.POST.get('score')
    sids = json.loads(request.POST.get('stuid'))

    # 如果評分為 0，清除該生成績紀錄
    if score == 0:
        Work.objects.filter(typing=2, user_id__in=sids, index=work_id, lesson_id=lesson).delete()
        return JsonResponse({'status':'ok'}, safe=False)

    # 取得舊評分紀錄，更新資料，如果找不到該筆紀錄就新增
    for sid in sids:
        try:
            work = Work.objects.get(typing=2, user_id=sid, index=work_id, lesson_id=lesson)
            work.publication_date = timezone.now()            
        except ObjectDoesNotExist:
            work = Work(typing=2, user_id=sid, index=work_id, lesson_id=lesson)
        except MultipleObjectsReturned:
            Work.objects.filter(typing=2, user_id=sid, index=work_id, lesson_id=lesson).delete()
            work = Work(typing=2, user_id=sid, index=work_id, lesson_id=lesson)
        work.score = request.POST.get('score')
        work.scorer = request.user.id
        work.save()
    return JsonResponse({'staus':'ok'}, safe=False)

# ck指定作業匯出
def work_word(request, lesson, classroom_id, index):
    if not is_teacher(request.user, classroom_id) and not is_assistant(request.user, classroom_id):
       return redirect("/")
    enroll_pool = [enroll for enroll in Enroll.objects.filter(classroom_id=classroom_id).order_by('seat')]
    student_ids = map(lambda a: a.student_id, enroll_pool)  
    works = Work.objects.filter(typing=0, lesson_id=lesson, user_id__in=student_ids).order_by("-id")
    classroom = Classroom.objects.get(id=classroom_id)
    lesson_list = lesson_list4
    #word
    document = Document()
    docx_title=u"指定作業-" + classroom.name + "-"+ index + "_" + lesson_list[int(index)][1]+ "-" + str(timezone.localtime(timezone.now()).date())+".docx"
    document.add_paragraph(request.user.first_name + u'的指定作業')
    document.add_paragraph(u'主題：'+ lesson_list[int(index)][1])		
    document.add_paragraph(u"班級：" + classroom.name)		

    for enroll in enroll_pool:
      enroll_works = filter(lambda w: w.user_id == enroll.student_id, works)
      work = enroll_works[0]
      run = document.add_paragraph().add_run(str(enroll.seat)+")"+enroll.student.first_name)
      font = run.font
      font.color.rgb = RGBColor(0xFA, 0x24, 0x00)
      if len(enroll_works)>0:
        p = document.add_paragraph(str(work.publication_date)[:19]+'\n')
        add_hyperlink(document, p, work.youtube, work.youtube)
        document.add_paragraph(work.memo)

    # Prepare document for download        
    f = StringIO.StringIO()
    document.save(f)
    length = f.tell()
    f.seek(0)
    response = HttpResponse(
      f.getvalue(),
      content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename={0}'.format(docx_title.encode('utf8')) 
    response['Content-Length'] = length
    return response

def add_hyperlink(document, paragraph, url, name):
    """
    Add a hyperlink to a paragraph.
    :param document: The Document being edited.
    :param paragraph: The Paragraph the hyperlink is being added to.
    :param url: The url to be added to the link.
    :param name: The text for the link to be displayed in the paragraph
    :return: None
    """

    part = document.part
    rId = part.relate_to(url, RT.HYPERLINK, is_external=True)

    init_hyper = OxmlElement('w:hyperlink')
    init_hyper.set(qn('r:id'), rId, )
    init_hyper.set(qn('w:history'), '1')

    new_run = OxmlElement('w:r')

    rPr = OxmlElement('w:rPr')

    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')

    rPr.append(rStyle)
    new_run.append(rPr)
    new_run.text = name
    init_hyper.append(new_run)

    r = paragraph.add_run()
    r._r.append(init_hyper)
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return None  
  
# 列出所有討論主題
class ForumListView(ListView):
    model = FWork
    context_object_name = 'forums'
    template_name = "teacher/forum_list.html"		
    paginate_by = 20
    def get_queryset(self):        
        fclasses = FClass.objects.filter(classroom_id=self.kwargs['classroom_id']).order_by("-publication_date", "-forum_id")
        forums = []
        for fclass in fclasses:
            forum = FWork.objects.get(id=fclass.forum_id)
            forums.append([forum, fclass])
        return forums
			
    def get_context_data(self, **kwargs):
        context = super(ForumListView, self).get_context_data(**kwargs)
        classroom = Classroom.objects.get(id=self.kwargs['classroom_id'])
        context['classroom'] = classroom
        return context	
        
#新增一個討論主題
class ForumCreateView(CreateView):
    model = FWork
    form_class = ForumForm
    template_name = "teacher/forum_form.html"
    def form_valid(self, form):
        self.object = form.save(commit=False)
        self.object.teacher_id = self.request.user.id
        self.object.classroom_id = self.kwargs['classroom_id']
        self.object.domains = self.request.POST.getlist('domains')
        self.object.levels = self.request.POST.getlist('levels')	        
        self.object.save()  
        classrooms = self.request.POST.getlist('classrooms')
        for classroom in classrooms:
          forum_class = FClass(forum_id=self.object.id, classroom_id=classroom)
          forum_class.save()
        
        return redirect("/teacher/forum/"+self.kwargs['classroom_id'])           
        
    def get_context_data(self, **kwargs):
        context = super(ForumCreateView, self).get_context_data(**kwargs)
        classroom_list = []
        classrooms = Classroom.objects.filter(teacher_id=self.request.user.id)
        for classroom in classrooms:
            classroom_list.append(classroom.id)
        assistants = Assistant.objects.filter(user_id=self.request.user.id)
        for assistant in assistants:
            if not assistant.classroom_id in classroom_list:
                classroom_list.append(assistant.classroom_id)
        classrooms = Classroom.objects.filter(id__in=classroom_list).order_by("-id")
        context['classrooms'] = classrooms
        context['classroom_id'] = int(self.kwargs['classroom_id'])
        context['classroom'] = Classroom.objects.get(id=self.kwargs['classroom_id'])
        return context	
  
        return redirect("/teacher/forum/"+self.kwargs['classroom_id'])        
	
def forum_categroy(request, classroom_id, forum_id):
    forum = FWork.objects.get(id=forum_id)
    domains = Domain.objects.all()
    levels = Level.objects.all()		
    if request.method == 'POST':
        form = ForumCategroyForm(request.POST)
        if form.is_valid():
            forum.domains = request.POST.getlist('domains')
            forum.levels = request.POST.getlist('levels')	
            forum.save()
            return redirect('/teacher/forum/'+classroom_id+'/#'+str(forum.id))
    else:
        form = CategroyForm(instance=forum)
        
    return render_to_response('teacher/forum_categroy_form.html',{'domains': domains, 'levels':levels, 'classroom_id': classroom_id, 'forum':forum}, context_instance=RequestContext(request))

	
# 列出所有討論主題
class ForumAllListView(ListView):
    model = FWork
    context_object_name = 'forums'
    template_name = "teacher/forum_all.html"		
    paginate_by = 20
		
    def get_queryset(self):
      # 年級
      if self.kwargs['categroy'] == "1":
        queryset = FWork.objects.filter(levels__contains=self.kwargs['categroy_id']).order_by("-id")
      # 學習領域
      elif self.kwargs['categroy'] == "2":
        queryset = FWork.objects.filter(domains__contains=self.kwargs['categroy_id']).order_by("-id")   
      else:
        queryset = FWork.objects.all().order_by("-id")
      if self.request.GET.get('account') != None:
        keyword = self.request.GET.get('account')
        users = User.objects.filter(Q(username__icontains=keyword) | Q(first_name__icontains=keyword)).order_by('-id')
        user_list = []
        for user in users:
            user_list.append(user.id)
        forums = queryset.filter(teacher_id__in=user_list)
        return forums
      else:				
        return queryset
			
    def get_context_data(self, **kwargs):
        context = super(ForumAllListView, self).get_context_data(**kwargs)
        context['categroy'] = self.kwargs['categroy']							
        context['categroy_id'] = self.kwargs['categroy_id']							
        context['levels'] = Level.objects.all()				
        context['domains'] = Domain.objects.all()
        return context	

# 展示討論素材
def forum_show(request, forum_id):
    forum = FWork.objects.get(id=forum_id)
    domains = Domain.objects.all()
    domain_dict = {}
    for domain in domains :
        key = domain.id
        domain_dict[key] = domain
    levels = Level.objects.all()	
    level_dict = {}
    for level in levels :
        key = level.id
        level_dict[key] = level
    contents = FContent.objects.filter(forum_id=forum_id)
    domains = []		
    if forum.domains:
        forum_domains = ast.literal_eval(forum.domains)
        for domain in forum_domains:
            key = int(domain)
            domains.append(domain_dict[key])
    levels = []						
    if forum.levels:
        forum_levels = ast.literal_eval(forum.levels)
        for level in forum_levels:
            key = int(level)			
            levels.append(level_dict[key])
    return render_to_response('teacher/forum_show.html',{'domains':domains, 'levels':levels, 'contents':contents, 'forum':forum}, context_instance=RequestContext(request))

		
# 列出某討論主題的班級
class ForumClassListView(ListView):
    model = FWork
    context_object_name = 'classrooms'
    template_name = "teacher/forum_class.html"		
    paginate_by = 20
	
    def get_queryset(self):        		
        fclass_dict = dict(((fclass.classroom_id, fclass) for fclass in FClass.objects.filter(forum_id=self.kwargs['forum_id'])))		
        classroom_list = []
        classroom_ids = []
        classrooms = Classroom.objects.filter(teacher_id=self.request.user.id).order_by("-id")
        for classroom in classrooms:
            if classroom.id in fclass_dict:
                classroom_list.append([classroom, True, fclass_dict[classroom.id].deadline, fclass_dict[classroom.id].deadline_date])
            else :
                classroom_list.append([classroom, False, False, timezone.now()])
            classroom_ids.append(classroom.id)
        assistants = Assistant.objects.filter(user_id=self.request.user.id)
        for assistant in assistants:
            classroom = Classroom.objects.get(id=assistant.classroom_id)
            if not classroom.id in classroom_ids:
                if classroom.id in fclass_dict:
                    classroom_list.append([classroom, True, fclass_dict[classroom.id].deadline, fclass_dict[classroom.id].deadline_date])
                else :
                    classroom_list.append([classroom, False, False, timezone.now()])
        return classroom_list
			
    def get_context_data(self, **kwargs):
        context = super(ForumClassListView, self).get_context_data(**kwargs)				
        fwork = FWork.objects.get(id=self.kwargs['forum_id'])
        context['fwork'] = fwork
        context['forum_id'] = self.kwargs['forum_id']
        return context	
	
# Ajax 開放班取、關閉班級
def forum_switch(request):
    forum_id = request.POST.get('forumid')
    classroom_id = request.POST.get('classroomid')		
    status = request.POST.get('status')
    try:
        fwork = FClass.objects.get(forum_id=forum_id, classroom_id=classroom_id)
        if status == 'false' :
    				fwork.delete()
    except ObjectDoesNotExist:
        if status == 'true':
            fwork = FClass(forum_id=forum_id, classroom_id=classroom_id)
            fwork.save()
    return JsonResponse({'status':status}, safe=False)        
	
# 列出某作業所有同學名單
def forum_class(request, classroom_id, work_id):
    enrolls = Enroll.objects.filter(classroom_id=classroom_id)
    classroom_name = Classroom.objects.get(id=classroom_id).name
    classmate_work = []
    scorer_name = ""
    for enroll in enrolls:
        try:    
            work = SWork.objects.get(student_id=enroll.student_id, index=work_id)
            if work.scorer > 0 :
                scorer = User.objects.get(id=work.scorer)
                scorer_name = scorer.first_name
            else :
                scorer_name = "1"
        except ObjectDoesNotExist:
            work = SWork(index=work_id, student_id=1)
        try:
            group_name = EnrollGroup.objects.get(id=enroll.group).name
        except ObjectDoesNotExist:
            group_name = "沒有組別"
        assistant = Assistant.objects.filter(classroom_id=classroom_id, student_id=enroll.student_id, lesson=work_id)
        if assistant.exists():
            classmate_work.append([enroll,work,1, scorer_name, group_name])
        else :
            classmate_work.append([enroll,work,0, scorer_name, group_name])   
    def getKey(custom):
        return custom[0].seat
	
    classmate_work = sorted(classmate_work, key=getKey)
   
    return render_to_response('teacher/twork_class.html',{'classmate_work': classmate_work, 'classroom_id':classroom_id, 'index': work_id}, context_instance=RequestContext(request))

# 列出所有討論主題素材
class ForumContentListView(ListView):
    model = FContent
    context_object_name = 'contents'
    template_name = "teacher/forum_content.html"		
    def get_queryset(self):
        queryset = FContent.objects.filter(forum_id=self.kwargs['forum_id']).order_by("-id")
        return queryset
			
    def get_context_data(self, **kwargs):
        context = super(ForumContentListView, self).get_context_data(**kwargs)
        fwork = FWork.objects.get(id=self.kwargs['forum_id'])
        fclasses = FClass.objects.filter(forum_id=self.kwargs['forum_id'])				
        context['fwork']= fwork
        context['forum_id'] = self.kwargs['forum_id']
        context['fclasses'] = fclasses
        return context	
			
#新增一個課程
class ForumContentCreateView(CreateView):
    model = FContent
    form_class = ForumContentForm
    template_name = "teacher/forum_content_form.html"
    def form_valid(self, form):
        self.object = form.save(commit=False)
        work = FContent(forum_id=self.object.forum_id)
        if self.object.types == 1:
            work.types = 1
            work.title = self.object.title
            work.link = self.object.link
        if self.object.types  == 2:
            work.types = 2					
            work.youtube = self.object.youtube
        if self.object.types  == 3:
            work.types = 3
            myfile = self.request.FILES['content_file']
            fs = FileSystemStorage()
            filename = uuid4().hex
            work.title = myfile.name
            work.filename = str(self.request.user.id)+"/"+filename
            fs.save("static/upload/"+str(self.request.user.id)+"/"+filename, myfile)
        if self.object.types  == 4:
            work.types = 4
        work.memo = self.object.memo
        work.save()         
  
        return redirect("/teacher/forum/content/"+self.kwargs['forum_id'])  

    def get_context_data(self, **kwargs):
        ctx = super(ForumContentCreateView, self).get_context_data(**kwargs)
        ctx['forum'] = FWork.objects.get(id=self.kwargs['forum_id'])
        return ctx

def forum_delete(request, forum_id, content_id):
    instance = FContent.objects.get(id=content_id)
    instance.delete()

    return redirect("/teacher/forum/content/"+forum_id)  
	
def forum_edit(request, forum_id, content_id):
    try:
        instance = FContent.objects.get(id=content_id)
    except:
        pass
    if request.method == 'POST':
            content_id = request.POST.get("id", "")
            try:
                content = FContent.objects.get(id=content_id)
            except ObjectDoesNotExist:
	              content = FContent(forum_id= request.POST.get("forum_id", ""), types=form.cleaned_data['types'])
            if content.types == 1:
                content.title = request.POST.get("title", "")
                content.link = request.POST.get("link", "")
            elif content.types == 2:
                content.youtube = request.POST.get("youtube", "")
            elif content.types == 3:
                myfile =  request.FILES.get("content_file", "")
                fs = FileSystemStorage()
                filename = uuid4().hex
                content.title = myfile.name
                content.filename = str(request.user.id)+"/"+filename
                fs.save("static/upload/"+str(request.user.id)+"/"+filename, myfile)
            content.memo = request.POST.get("memo", "")
            content.save()
            return redirect('/teacher/forum/content/'+forum_id)   
    return render_to_response('teacher/forum_edit.html',{'content': instance, 'forum_id':forum_id, 'content_id':content_id}, context_instance=RequestContext(request))		
	
def forum_download(request, content_id):
    content = FContent.objects.get(id=content_id)
    filename = content.title
    BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))		
    download =  BASE_DIR + "/static/upload/" + content.filename
    wrapper = FileWrapper(file( download, "r" ))
    response = HttpResponse(wrapper, content_type = 'application/force-download')
    #response = HttpResponse(content_type='application/force-download')
    response['Content-Disposition'] = 'attachment; filename={0}'.format(filename.encode('utf8'))
    # It's usually a good idea to set the 'Content-Length' header too.
    # You can also set any other required headers: Cache-Control, etc.
    return response
    #return render_to_response('student/download.html', {'download':download})
		
class ForumEditUpdateView(UpdateView):
    model = FWork
    fields = ['title']
    template_name = 'form.html'
    #success_url = '/teacher/forum/domain/'
    def get_success_url(self):
        succ_url =  '/teacher/forum/'+self.kwargs['classroom_id']
        return succ_url
	
def forum_export(request, classroom_id, forum_id):
	if not is_teacher(request.user, classroom_id):
		return redirect("/")
	classroom = Classroom.objects.get(id=classroom_id)
	try:
		fwork = FWork.objects.get(id=forum_id)
		enrolls = Enroll.objects.filter(classroom_id=classroom_id)
		datas = []
		contents = FContent.objects.filter(forum_id=forum_id).order_by("-id")
		fwork = FWork.objects.get(id=forum_id)
		works_pool = SFWork.objects.filter(index=forum_id).order_by("-id")
		reply_pool = SFReply.objects.filter(index=forum_id).order_by("-id")	
		file_pool = SFContent.objects.filter(index=forum_id, visible=True).order_by("-id")	
		for enroll in enrolls:
			works = list(filter(lambda w: w.student_id==enroll.student_id, works_pool))
			if len(works)>0:
				replys = list(filter(lambda w: w.work_id==works[0].id, reply_pool))
			else:
				replys = []
			files = list(filter(lambda w: w.student_id==enroll.student_id, file_pool))
			if enroll.seat > 0:
				datas.append([enroll, works, replys, files])
		def getKey(custom):
			return -custom[0].seat
		datas = sorted(datas, key=getKey, reverse=True)	
		#word
		document = Document()
		docx_title=u"討論區-" + classroom.name + "-"+ str(timezone.localtime(timezone.now()).date())+".docx"
		document.add_paragraph(request.user.first_name + u'的討論區作業')
		document.add_paragraph(u'主題：'+fwork.title)		
		document.add_paragraph(u"班級：" + classroom.name)		
		
		for enroll, works, replys, files in datas:
			user = User.objects.get(id=enroll.student_id)
			run = document.add_paragraph().add_run(str(enroll.seat)+")"+user.first_name)
			font = run.font
			font.color.rgb = RGBColor(0xFA, 0x24, 0x00)
			if len(works)>0:
				#p = document.add_paragraph(str(works[0].publication_date)[:19]+'\n'+works[0].memo)
				p = document.add_paragraph(str(localtime(works[0].publication_date))[:19]+'\n')
				# 將 memo 以時間標記為切割點，切分為一堆 tokens
				tokens = re.split('(\[m_\d+#\d+:\d+:\d+\])', works[0].memo)
				# 依續比對 token 格式
				for token in tokens:
					m = re.match('\[m_(\d+)#(\d+):(\d+):(\d+)\]', token)
					if m: # 若為時間標記，則插入連結
						vid = filter(lambda material: material.id == int(m.group(1)), contents)[0]
						add_hyperlink(document, p, vid.youtube+"&t="+m.group(2)+"h"+m.group(3)+"m"+m.group(4)+"s", "["+m.group(2)+":"+m.group(3)+":"+m.group(4)+"]")
					else: # 以一般文字插入
						p.add_run(token)
			if len(replys)>0:
				for reply in replys:
					user = User.objects.get(id=reply.user_id)
					run = document.add_paragraph().add_run(user.first_name+u'>'+str(localtime(reply.publication_date))[:19]+u'>留言:\n'+reply.memo)
					font = run.font
					font.color.rgb = RGBColor(0x42, 0x24, 0xE9)		
			if len(files)>0:
				for file in files:
					if file.visible:
						if file.title[-3:].upper() == "PNG" or file.title[-3:].upper() == "JPG":                     
						    filename = 'static/upload/'+file.filename
						    if os.path.exists(filename):
						        im = Image.open(filename)
						        im.save('static/upload/file.'+ file.title[-3:]')				
						        document.add_picture('static/upload/file.'+file.title[-3:],width=Inches(6.0))  
						else:
							p = document.add_paragraph()
							full_url = request.build_absolute_uri()
							index = full_url.find("/",9)
							url = full_url[:index] + "/student/forum/download/" + str(file.id) 
							add_hyperlink(document, p, url, file.title)
		# Prepare document for download        
		f = io.BytesIO()
		document.save(f)
		length = f.tell()
		f.seek(0)
		response = HttpResponse(
			f.getvalue(),
			content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
		)

		filename_header = filename_browser(request, docx_title)
		response['Content-Disposition'] = 'attachment; ' + filename_header
		response['Content-Length'] = length
		return response

	except ObjectDoesNotExist:
		pass
	return True		


def add_hyperlink(document, paragraph, url, name):
    """
    Add a hyperlink to a paragraph.
    :param document: The Document being edited.
    :param paragraph: The Paragraph the hyperlink is being added to.
    :param url: The url to be added to the link.
    :param name: The text for the link to be displayed in the paragraph
    :return: None
    """

    part = document.part
    rId = part.relate_to(url, RT.HYPERLINK, is_external=True)

    init_hyper = OxmlElement('w:hyperlink')
    init_hyper.set(qn('r:id'), rId, )
    init_hyper.set(qn('w:history'), '1')

    new_run = OxmlElement('w:r')

    rPr = OxmlElement('w:rPr')

    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')

    rPr.append(rStyle)
    new_run.append(rPr)
    new_run.text = name
    init_hyper.append(new_run)

    r = paragraph.add_run()
    r._r.append(init_hyper)
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return None

def forum_grade(request, classroom_id, action):
	classroom = Classroom.objects.get(id=classroom_id)
	forum_ids = []
	forums = []
	fclasses = FClass.objects.filter(classroom_id=classroom_id).order_by("publication_date", "forum_id")
	for fclass in fclasses:
		forum_ids.append(fclass.forum_id)
		forum = FWork.objects.get(id=fclass.forum_id)
		forums.append(forum.title)
	enrolls = Enroll.objects.filter(classroom_id=classroom_id).order_by("seat")
	datas = {}
	for enroll in enrolls:
			sfworks = SFWork.objects.filter(index__in=forum_ids, student_id=enroll.student_id).order_by("id")
			if len(sfworks) > 0:
				for fclass in fclasses:
						works = filter(lambda w: w.index==fclass.forum_id, sfworks)
						if enroll.student_id in datas:
							if len(works) > 0 :
								datas[enroll.student_id].append(works[0])
							else :
								datas[enroll.student_id].append(SFWork())
						else:
							if len(works) > 0:
								datas[enroll.student_id] = [works[0]]
							else :
								datas[enroll.student_id] = [SFWork()]
			else :
				datas[enroll.student_id] = [SFWork()]
	results = []
	for enroll in enrolls:
		student_name = User.objects.get(id=enroll.student_id).first_name
		results.append([enroll, student_name, datas[enroll.student_id]])
	
	#下載Excel
	if action == "1":
		classroom = Classroom.objects.get(id=classroom_id)       
		output = StringIO.StringIO()
		workbook = xlsxwriter.Workbook(output)    
		worksheet = workbook.add_worksheet(classroom.name)
		date_format = workbook.add_format({'num_format': 'yy/mm/dd'})
		
		row = 1
		worksheet.write(row, 1, u'座號')
		worksheet.write(row, 2, u'姓名')
		index = 3
		for forum in forums:
			worksheet.write(row, index, forum)
			index += 1
		
		row += 1
		index = 3
		for fclass in fclasses:
			worksheet.write(row, index, datetime.strptime(str(fclass.publication_date)[:19],'%Y-%m-%d %H:%M:%S'), date_format)
			index += 1			

		for enroll, student_name, works in results:
			row += 1
			worksheet.write(row, 1, enroll.seat)
			worksheet.write(row, 2, student_name)
			index = 3
			for work in works:
				if work.id:
					worksheet.write(row, index, work.score)
				else:
					worksheet.write(row, index, '')
				index +=1 

		workbook.close()
		# xlsx_data contains the Excel file
		response = HttpResponse(content_type='application/vnd.ms-excel')
		filename = classroom.name + '-' + str(localtime(timezone.now()).date()) + '.xlsx'
		response['Content-Disposition'] = 'attachment; filename={0}'.format(filename.encode('utf8'))
		xlsx_data = output.getvalue()
		response.write(xlsx_data)
		return response
	else :
		return render(request, 'teacher/forum_grade.html',{'results':results, 'forums':forums, 'classroom_id':classroom_id, 'fclasses':fclasses})

def forum_deadline(request, classroom_id, forum_id):
    forum = FWork.objects.get(id=forum_id)
    classroom = Classroom.objects.get(id=classroom_id)
    if request.method == 'POST':
        form = CategroyForm(request.POST)
        if form.is_valid():
            forum.domains = request.POST.getlist('domains')
            forum.levels = request.POST.getlist('levels')	
            forum.save()
            return redirect('/teacher/forum/'+classroom_id)
    else:
        fclass = FClass.objects.get(classroom_id=classroom_id, forum_id=forum_id)
        form = ForumDeadlineForm(instance=fclass)
        fclasses = FClass.objects.filter(forum_id=forum_id).order_by("-id")
    return render(request, 'teacher/forum_deadline_form.html',{'fclasses':fclasses, 'fclass':fclass, 'forum':forum, 'classroom':classroom})

	
# Ajax 設定期限、取消期限
def forum_deadline_set(request):
    forum_id = request.POST.get('forumid')
    classroom_id = request.POST.get('classroomid')		
    status = request.POST.get('status')
    try:
        fclass = FClass.objects.get(forum_id=forum_id, classroom_id=classroom_id)
    except ObjectDoesNotExist:
        fclass = Fclass(forum_id=forum_id, classroom_id=classroom_id)
    if status == 'True':
        fclass.deadline = True
    else :
        fclass.deadline = False
    fclass.save()
    return JsonResponse({'status':status}, safe=False)        

# Ajax 設定期限日期
def forum_deadline_date(request):
    forum_id = request.POST.get('forumid')
    classroom_id = request.POST.get('classroomid')		
    deadline_date = request.POST.get('deadlinedate')
    try:
        fclass = FClass.objects.get(forum_id=forum_id, classroom_id=classroom_id)
    except ObjectDoesNotExist:
        fclass = FClass(forum_id=forum_id, classroom_id=classroom_id)
    #fclass.deadline_date = deadline_date.strftime('%d/%m/%Y')
    fclass.deadline_date = datetime.strptime(deadline_date, '%Y %B %d - %H:%M')
    fclass.save()
    return JsonResponse({'status':deadline_date}, safe=False)     
	
class ForumPublishReject(RedirectView):
    def get_redirect_url(self, *args, **kwargs):
      index = self.kwargs['index']
      classroom_id = self.kwargs['classroom_id']
      user_id = self.kwargs['user_id']
      try:
          fwork = FWork.objects.get(id=index)
          works = SFWork.objects.filter(index=index, student_id=user_id).order_by("-id")
          work = works[0]
          work.publish = False
          work.save()
          update_avatar(user_id, 4, -2)
          # History
          history = PointHistory(user_id=user_id, kind=4, message=u'-2分--退回討論區作業<'+fwork.title+'>', url='/student/forum/memo/'+str(classroom_id)+'/'+str(index)+'/0')
          history.save()								
      except ObjectDoesNotExist:
            pass
      return "/student/forum/memo/"+str(classroom_id)+"/"+str(index)+"/0"
	
	
# 影片觀看時間統計
class EventVideoView(ListView):
    context_object_name = 'events'
    #paginate_by = 50
    template_name = 'teacher/event_video.html'

    def get_queryset(self):    
				enrolls = Enroll.objects.filter(classroom_id=self.kwargs['classroom_id'], seat__gt=0).order_by("seat")
				events = []
				for enroll in enrolls: 
						videos = VideoLogHelper().getLogByUserid(enroll.student_id,self.kwargs['work_id'])
						length = 0
						for video in videos: 
										length += video['length']
						events.append([enroll, length/60.0])
				return events
			
    def get_context_data(self, **kwargs):
        context = super(EventVideoView, self).get_context_data(**kwargs)
        classroom = Classroom.objects.get(id=self.kwargs['classroom_id'])
        context['length'] = FContent.objects.get(id=self.kwargs['work_id']).youtube_length / 60.0
        context['content_id'] = self.kwargs['work_id']
        context['classroom'] = classroom
        enrolls = Enroll.objects.filter(classroom_id=classroom.id)
        context['height'] = 100 + enrolls.count() * 40
        return context
			
# 記錄影片長度
def video_length(request):
    content_id = request.POST.get('content_id')
    length = request.POST.get('length')
    fcontent = FContent.objects.get(id=content_id)
    fcontent.youtube_length = length
    fcontent.save()
    return JsonResponse({'status':'ok'}, safe=False)	
	
# 影片記錄條
class VideoListView(ListView):
    context_object_name = 'videos'
    template_name = 'teacher/event_video_user.html'
    
    def get_queryset(self):
				videos = VideoLogHelper().getLogByUserid(self.kwargs['user_id'],self.kwargs['content_id'])        
				return videos
        
    def get_context_data(self, **kwargs):
        context = super(VideoListView, self).get_context_data(**kwargs)
        content = FContent.objects.get(id=self.kwargs['content_id'])
        context['user_id'] = self.kwargs['user_id']
        context['content'] = content
        context['length'] = content.youtube_length
        return context  

    # 限本班任課教師或助教     
    def render_to_response(self, context):
        if not is_teacher(self.request.user ,self.kwargs['classroom_id']):
            if not is_assistant(self.request.user, self.kwargs['classroom_id'] ):
                  return redirect('/')
        return super(VideoListView, self).render_to_response(context) 
      
# 列出所有作業小老師
def assistant_group(request, typing, classroom_id):
        # 限本班任課教師
        if not is_teacher(request.user, classroom_id) and not is_assistant(request.user, classroom_id):
            return redirect("/")    
        classroom = Classroom.objects.get(id=classroom_id)
        lesson = Classroom.objects.get(id=classroom_id).lesson        
        if typing == "0":
            if lesson == 1:
                lesson_name = lesson_list1
            elif lesson == 2:
                lesson_name = lesson_list2
            elif lesson == 3:
                lesson_name = lesson_list3
            elif lesson == 4:
                lesson_name = lesson_list4
            elif lesson == 5:
                lesson_name = lesson_list2
            elif lesson == 8:
                lesson_name = lesson_list5       
            else:
                lesson_name = lesson_list1
        elif typing == "1":
            lesson_name = TWork.objects.get(classroom_id=classroom_id).title        
        groups = range(classroom.group_number)				
        enroll_pool = [enroll for enroll in Enroll.objects.filter(classroom_id=classroom_id).order_by('seat')]
        student_ids = map(lambda a: a.student_id, enroll_pool)
        work_pool = Work.objects.filter(user_id__in=student_ids, lesson_id=classroom.lesson)
        user_pool = [user for user in User.objects.filter(id__in=work_pool.values('scorer'))]
        assistant_pool = [assistant for assistant in WorkAssistant.objects.filter(classroom_id=classroom_id, typing=typing, lesson_id=lesson)]				
        lessons = []		
        index = 1
        for assignment in lesson_name:
                student_groups = []													
                for group in groups:
                    members = filter(lambda u: u.group == group, enroll_pool)
                    group_assistants = []
                    works = []
                    scorer_name = ""
                    for member in members:
                        work = filter(lambda w: w.index == index and w.user_id == member.student_id, work_pool)
                        if work:
                            work = work[0]
                            scorer = filter(lambda u: u.id == work.scorer, user_pool)
                            scorer_name = scorer[0].first_name if scorer else 'X'
                        else:
                            work = Work(index=assignment[2], user_id=1, score=-2)
                        works.append([member, work.score, scorer_name, work.memo])
                        assistant = filter(lambda a: a.student_id == member.student_id and a.index == index, assistant_pool)
                        if assistant:
                            group_assistants.append(member)
                    student_groups.append([group, works, group_assistants])                    
                lessons.append([assignment, student_groups])
                index = index + 1
        return render(request, 'teacher/assistant_group.html', {'lessons':lessons,'classroom':classroom})

# 科學運算現象描述問題
class Science1QuestionListView(ListView):
    model = Science1Question
    context_object_name = 'questions'
    template_name = 'teacher/question_list.html'

    def get_queryset(self):
        queryset = Science1Question.objects.filter(work_id=self.kwargs['work_id']).order_by("-id")
        return queryset

    def get_context_data(self, **kwargs):
        context = super(Science1QuestionListView, self).get_context_data(**kwargs)
        context['classroom'] = Classroom.objects.get(id=self.kwargs['classroom_id'])
        context['lesson'] = self.kwargs['lesson']
        context['work_id'] = self.kwargs['work_id']		
        return context
		
# 科學運算現象描述問題
class Science1QuestionAnswerView(ListView):
    model = Science1Work
    context_object_name = 'works'
    template_name = 'teacher/question_answer.html'

    def get_queryset(self):
        enroll_pool = [enroll for enroll in Enroll.objects.filter(classroom_id=self.kwargs['classroom_id'], seat__gt=0).order_by('seat')]
        student_ids = map(lambda a: a.student_id, enroll_pool)
        work_pool = Science1Work.objects.filter(student_id__in=student_ids, question_id=self.kwargs['q_id'])
        work_ids = map(lambda a: a.id, work_pool)
        content_pool = Science1Content.objects.filter(work_id__in=work_ids)
        queryset = []
        for enroll in enroll_pool:
            works = filter(lambda w: w.student_id==enroll.student_id, work_pool)
            if works:
			    contents = filter(lambda w: w.work_id==works[0].id, content_pool)
            else :
                contents = []
            queryset.append([enroll, contents])
        return queryset

    def get_context_data(self, **kwargs):
        context = super(Science1QuestionAnswerView, self).get_context_data(**kwargs)
        context['classroom'] = Classroom.objects.get(id=self.kwargs['classroom_id'])
        context['lesson'] = self.kwargs['lesson']
        context['work_id'] = self.kwargs['work_id']
        context['question'] = Science1Question.objects.get(id=self.kwargs['q_id'])		
        return context		

#新增一個問題
class Science1QuestionCreateView(CreateView):
    model = Science1Question
    form_class = QuestionForm
    template_name = 'form.html'
    def form_valid(self, form):
        self.object = form.save(commit=False)
        self.object.work_id = self.kwargs['work_id']
        self.object.save()
        return redirect("/teacher/work2/question/"+self.kwargs['lesson']+"/"+self.kwargs['classroom_id']+"/"+self.kwargs['work_id'])

def work2_science(request, classroom_id, index, user_id):
    contents1 = [[]]
    works_pool = Science1Work.objects.filter(student_id=user_id, index=index).order_by("-id")
    questions = Science1Question.objects.filter(work_id=index)			
    for question in questions:
        works = filter(lambda w: w.question_id==question.id, works_pool)
        if len(works) > 0:
            contents = Science1Content.objects.filter(work_id=works[0].id).order_by("id")
            if len(contents)>0:
                contents1.append([contents])
            else:
                contents1.append([[]])						
        else:
            contents1.append([[]])
    try:
        work4 = Science4Work.objects.get(student_id=user_id, index=index)
    except ObjectDoesNotExist:
        work4 = Science4Work(student_id=user_id, index=index)                        
    except MultipleObjectsReturned:
        works4 = Science4Work.objects.filter(student_id=user_id, index=index).order_by("-id")
        work4 = works[0]
    contents4 = Science4Content.objects.filter(work_id=work4.id).order_by("id")            
    works3 = Science3Work.objects.filter(student_id=user_id, index=index).order_by("-id")
    if works3.exists():
        work3 = works3[0]
    else :
        work3 = Science3Work(student_id=user_id, index=index)
    data1 = Science2Data.objects.filter(index=index, student_id=user_id, types=0).order_by("id")
    data2 = Science2Data.objects.filter(index=index, student_id=user_id, types=1).order_by("id")
    data3 = Science2Data.objects.filter(index=index, student_id=user_id, types=2).order_by("id")				
    questions = Science1Question.objects.filter(work_id=index)	
    return render(request, 'teacher/work2_science.html', {'user_id':user_id, 'questions':questions, 'data1':data1, 'data2':data2, 'data3':data3, 'index':index, 'contents1':contents1, 'contents4':contents4, 'work3':work3})                  
			
# 查閱全班測驗卷成績
def exam_list(request, classroom_id):
        # 限本班任課教師
        if not is_teacher(request.user, classroom_id):
            return redirect("homepage")
        enrolls = Enroll.objects.filter(classroom_id=classroom_id).order_by("seat")
        classroom_name=""
        enroll_exam = []		
        for enroll in enrolls:
            classroom_name = enroll.classroom.name
            exam_list = []
            for exam_id in range(3):
                exams = Exam.objects.filter(student_id=enroll.student_id, exam_id=exam_id+1)
                total = 0
                times = 0
                for exam in exams:
                    total += exam.score
                    times += 1
                exam_list.append(total)
                exam_list.append(times)
            enroll_exam.append([enroll, exam_list])
        return render(request, 'teacher/exam_list.html', {'classroom_id':classroom_id, 'classroom_name':classroom_name, 'enroll_exam':enroll_exam})

# 查詢某項測驗的所有資料
def exam_detail(request, classroom_id, student_id, exam_id):
        # 限本班任課教師
        if not is_teacher(request.user, classroom_id):
            return redirect("homepage")
        exams = Exam.objects.filter(student_id=student_id, exam_id=exam_id)
        enroll = Enroll.objects.get(classroom_id=classroom_id, student_id=student_id)  
        return render(request, 'teacher/exam_detail.html', {'exams': exams, 'enroll':enroll})	

class GroupUpdate(UpdateView):
    model = Classroom
    form_class = GroupForm	
    template_name = 'form.html'
			
    def get_success_url(self):
        succ_url =  '/student/group/panel/'+str(self.kwargs['pk'])
        return succ_url
			
    def form_valid(self, form):
        classroom = Classroom.objects.get(id=self.kwargs['pk'])
        if is_teacher(self.request.user, classroom.id) or is_assistant(self.request.user, classroom.id):
            form.save()
        return HttpResponseRedirect(self.get_success_url())

@user_passes_test(not_in_teacher_group, login_url='/')
def group_assign(request, classroom_id):
    if not is_teacher(request.user, classroom_id):
        return redirect("/")
    classroom = Classroom.objects.get(id=classroom_id)
    GroupModelFormset = modelformset_factory(Enroll, fields=['seat','student_id', 'group'],extra=0)	
    if request.method == 'POST':
        formset = GroupModelFormset(request.POST)
        if formset.is_valid():
            for form in formset:
                form.save()
            return redirect('/student/group/panel/'+str(classroom_id))
        else:
            return redirect("/")
    else:
        formset = GroupModelFormset(queryset=Enroll.objects.filter(classroom_id=classroom_id).order_by("seat"))		
    return render(request, 'teacher/group_assign.html', {'formset': formset, 'group_numbers':range(classroom.group_number)})		
		
class GroupUpdate2(UpdateView):
    model = Classroom
    form_class = GroupForm2	
    template_name = 'form.html'
			
    def get_success_url(self):
        succ_url =  '/student/group/panel/'+str(self.kwargs['pk'])
        return succ_url
			
    def form_valid(self, form):
        classroom = Classroom.objects.get(id=self.kwargs['pk'])
        if is_teacher(self.request.user, classroom.id) or is_assistant(self.request.user, classroom.id):
            form.save()
        return HttpResponseRedirect(self.get_success_url())
			
# 分組
@user_passes_test(not_in_teacher_group, login_url='/')
def make(request, classroom_id, action):
    if not is_teacher(request.user, classroom_id):
        return redirect("/")
    classroom = Classroom.objects.get(id=classroom_id)
    if action == "1":            
        classroom.group_open = True   
    else : 
        classroom.group_open = False
    classroom.save()      
    return redirect("/student/group/panel/"+str(classroom.id))

	
# 教師可以查看所有帳號
class StudentJoinView(ListView):
    context_object_name = 'users'
    paginate_by = 40
    template_name = 'teacher/student_join.html'

    def get_queryset(self):
        enrolls = Enroll.objects.filter(classroom_id=self.kwargs['classroom_id'])	
        user_ids = [enroll.student_id for enroll in enrolls]
        username = username__icontains=self.request.user.username+"_"
        if self.request.GET.get('account') != None:
            keyword = self.request.GET.get('account')
            queryset = User.objects.filter((~Q(id__in=user_ids)) & (Q(username__icontains=username+keyword) | (Q(first_name__icontains=keyword)) & Q(username__icontains=username))).order_by('-id')
        else :
            queryset = User.objects.filter(~Q(id__in=user_ids) & Q(username__icontains=username)).order_by('-id')
        return queryset

    def get_context_data(self, **kwargs):
        context = super(StudentJoinView, self).get_context_data(**kwargs)
        account = self.request.GET.get('account')
        context.update({'account': account})
        context['classroom_id']=self.kwargs['classroom_id']
        return context


#加選學生
class StudentEnrollView(RedirectView):

    def get(self, request, *args, **kwargs):
        classroom_id = self.kwargs['classroom_id'] 
        classroom = Classroom.objects.get(id=classroom_id)
        students = self.request.POST.getlist('student')
        jsonDec = json.decoder.JSONDecoder()	
        classroom_list = []		
        for student in students:
            student_id = student.split(":")[0]
            student_seat = student.split(":")[1] 
            enroll = Enroll(student_id=student_id, classroom_id=classroom_id, seat=student_seat)
            enroll.save()
            user = User.objects.get(id=student_id)
            profile = Profile.objects.get(user=user)
            if len(profile.classroom) > 0 :		
                classroom_list = jsonDec.decode(profile.classroom)
            classroom_list.append(str(self.kwargs['classroom_id']))
            profile.classroom = json.dumps(classroom_list)
            profile.save()			
            messages = Message.objects.filter(author_id=classroom.teacher_id, classroom_id=classroom_id, type=1)
            for message in messages:
                try:
                    messagepoll = MessagePoll.objects.get(message_type=1, message_id=message.id, reader_id=student_id, classroom_id=classroom_id)
                except ObjectDoesNotExist:
                    messagepoll = MessagePoll(message_type=1, message_id=message.id, reader_id=student_id, classroom_id=classroom_id)
                    messagepoll.save()	
                except MultipleObjectsReturned:
                    pass				
        return super(StudentEnrollView, self).get(self, request, *args, **kwargs)        
        
    def get_redirect_url(self, *args, **kwargs):
        #TaxRate.objects.get(id=int(kwargs['pk'])).delete()   
        return '/student/classmate/'+ str(self.kwargs['classroom_id'])
	
@login_required
@user_passes_test(not_in_teacher_group, login_url='/')
def work_ckexcel(request, classroom_id):
    typing = "0"
    lesson = "4"
    # 限本班任課教師
    if not is_teacher(request.user, classroom_id) and not is_assistant(request.user, classroom_id):
        return redirect("/")
    enrolls = Enroll.objects.filter(classroom_id=classroom_id).order_by('seat')
    user_ids = [enroll.student_id for enroll in enrolls]
    work_pool = Work.objects.filter(typing=typing, user_id__in=user_ids, lesson_id=lesson).order_by('id')
    data = []
    for enroll in enrolls:
      enroll_score = []
      total = 0
      stu_works = filter(lambda w: w.user_id == enroll.student_id, work_pool)
      if typing == "0":
        lesson_list = lesson_list4
      elif typing == "1":
        lesson_list = TWork.objects.filter(classroom_id=classroom_id)
      else :
        lesson_list = CWork.objects.filter(classroom_id=classroom_id)
      memo = ""
      grade = 0
      for index, assignment in enumerate(lesson_list4):
            if typing == "0": 
                works = list(filter(lambda w: w.index == int(index)+1, stu_works))
            else :
                works = filter(lambda w: w.index == assignment.id, stu_works)
            works_count = len(works)
            if works_count == 0:
                enroll_score.append(["X", index, Work()])
                if typing == "0" or typing == "1":
                    if not lesson == "4":
                        total += 60
            else:
                work = works[-1]
                enroll_score.append([work.score, index, work])
                if work.score == -2:
                    if typing == "0" or typing == "1":
                          if not lesson == "4":
                              total += 80
                else:
                    total += work.score

            if typing == "0":
                if lesson == "1":
                    memo = [enroll.score_memo1, enroll.score_memo2, enroll.score_memo3, enroll.score_memo4][int(unit)-1]
                elif lesson == "2":
                    memo = enroll.score_memo_vphysics
                elif lesson == "3":
                    memo = enroll.score_memo_euler
                elif lesson == "4":
                    memo = enroll.score_memo_vphysics2
                elif lesson == "5":
                    memo = enroll.score_memo_vphysics3
                elif lesson == "6":
                    memo = enroll.score_memo_microbit 
                elif lesson == "7":
                    memo = enroll.score_memo_pandas                         
                elif lesson == "8":
                    memo = enroll.score_memo_django
            elif typing == "1":
                memo = enroll.score_memo_custom
            if typing == "2":
                grade = total
            else :
                grade = int(total / len(lesson_list) * 0.6 + memo * 0.4)
      data.append([enroll, enroll_score, memo, grade])
                
    classroom = Classroom.objects.get(id=classroom_id)       
    output = StringIO.StringIO()
    workbook = xlsxwriter.Workbook(output)    
    worksheet = workbook.add_worksheet(classroom.name)
    date_format = workbook.add_format({'num_format': 'yy/mm/dd'})

    row = 1
    worksheet.write(row, 1, u'座號')
    worksheet.write(row, 2, u'姓名')
    worksheet.write(row, 3, u'成績')        
    worksheet.write(row, 4, u'心得')         
    index = 5
    for assignment in lesson_list:
        if typing == "0":
	          worksheet.write(row, index, assignment[1])
        else :
	          worksheet.write(row, index, assignment.title)              
        index += 1

    index = 5
    if not typing == "0":
        row += 1
        for assignment in lesson_list:            
            worksheet.write(row, index, datetime.strptime(str(assignment.time)[:19],'%Y-%m-%d %H:%M:%S'), date_format)
            index += 1			

    for enroll, enroll_score, memo, grade in data:
      row += 1
      worksheet.write(row, 1, enroll.seat)
      worksheet.write(row, 2, enroll.student.first_name)
      worksheet.write(row, 3, grade)         
      worksheet.write(row, 4, memo)     
      index = 5
      for score, index2, work in enroll_score:
          if score == -2 :
              worksheet.write(row, index, "V")
          else :
              worksheet.write(row, index, score)
          index +=1 

    for index3, assignment in lesson_list4:
        worksheet = workbook.add_worksheet(index3+"_"+assignment)
        row = 1
        for enroll, enroll_score, memo, grade in data:
                worksheet.write(row, 1, enroll.seat)
                worksheet.write(row, 2, enroll.student.first_name)
                if enroll_score[int(index3)-1][0] != "X":
                    work = enroll_score[int(index3)-1][2]
                    worksheet.write(row, 3, work.youtube)
                    worksheet.write(row, 4, enroll_score[int(index3)-1][0])
                    worksheet.write(row, 5, work.memo)
                else :
                    worksheet.write(row, 3, "未繳交")
                row +=1

    workbook.close()
    # xlsx_data contains the Excel file
    response = HttpResponse(content_type='application/vnd.ms-excel')
    if typing == "0":
        type_name = "指定作業"
    elif typing == "1":
        type_name = "自訂作業"
    else:
        type_name = "檢核作業"        
    filename = classroom.name + '-' + type_name + "-" + str(localtime(timezone.now()).date()) + '.xlsx'
    response['Content-Disposition'] = 'attachment; filename={0}'.format(filename.encode('utf8'))
    xlsx_data = output.getvalue()
    response.write(xlsx_data)
    return response
